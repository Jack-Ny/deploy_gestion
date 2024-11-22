from collections import defaultdict
from django.db.models import Sum, Q, Count, Case, When, IntegerField, F
from django.db.models.functions import Coalesce
from io import BytesIO
import json
import os
from account.models import User
import openpyxl
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django import template
from django.core.exceptions import ObjectDoesNotExist
from django.db import transaction
from django.http import FileResponse, HttpResponse, JsonResponse
from django.shortcuts import redirect, render
from django.contrib import messages
from django.shortcuts import get_object_or_404
import xlsxwriter
from django.template.loader import render_to_string
from projetGestion import settings
from .models import *
from django.contrib.auth.decorators import login_required
from django.contrib.auth import get_user_model
from .decorators import responsable_required, responsable_charger_projet_required, all_user_required
from xhtml2pdf import pisa
from django.template.loader import get_template
from docx import Document
from openpyxl import Workbook
from django.contrib.auth.forms import PasswordChangeForm
from django.contrib.auth import update_session_auth_hash


# fonctions utilitaires
def safe_sum(values):
    return sum(value if value is not None else 0 for value in values)

def safe_join(values, separator=', ', default_value='pas renseigné'):
    return separator.join(value if value is not None else default_value for value in values)


def dashboard_view(request):
    total_service = Projet.objects.count()
    total_planifier = Activite.objects.count()
    total_realisation = Realisation.objects.count()
    total_user = User.objects.count()
    # Pour le graphique de réalisation
    projets = Projet.objects.annotate(
        activites_count=Coalesce(Count('id_projet', distinct=True), 0),
        realisations_count=Coalesce(Count('id_projet_plus', distinct=True), 0),
        taux_realisation=Case(
            When(activites_count__gt=0, 
                 then=F('realisations_count') * 100.0 / F('activites_count')),
            default=0,
            output_field=IntegerField(),
        )
    ).filter(activites_count__gt=0)

    projets_data = []
    for projet in projets:
        projets_data.append({
            'nom': projet.nom,
            'taux': float(projet.taux_realisation)
        })

    # Pour le graphique des budgets
    infos_budgets = InfosSpecific.objects.filter(
        budget__gt=0
    ).values('nom', 'budget', 'depense_globale')

    budgets_data = []
    for info in infos_budgets:
        budgets_data.append({
            'nom': info['nom'],
            'budget_total': float(info['budget'] or 0),
            'depenses_total': float(info['depense_globale'] or 0)
        })

    context = {
        'projets_data': projets_data,
        'budgets_data': budgets_data,
        'total_service' : total_service,
        'total_planifier': total_planifier,
        'total_realisation': total_realisation,
        'total_user': total_user
    }
    
    return render(request, 'service/dashboard.html', context)


@login_required(login_url='/login/')
def get_partenaires(request):
    #
    search_term = request.GET.get('term', '')

    # filtrer les partenaires qui commencent par search_term
    if search_term:
        partenaires = ListPartenaire.objects.filter(
            nom__istartswith = search_term
        ).values('id', 'nom')
    else:
        partenaires = ListPartenaire.objects.all().values('id', 'nom')

    # formater le resultats
    results = [{'id': p['nom'], 'text': p['nom']} for p in partenaires]

    return JsonResponse({
        'results': results,
        'pagination': { 'more': False }
    })

# generer un PDF
def generate_pdf(request):
    # Récupérer les données des projets
    projets = Projet.objects.all()
    synthese = []

    for projet in projets:
        infos_specific = InfosSpecific.objects.filter(id_projet=projet)
        depenses = Depense.objects.filter(id_projet=projet)
        realisations = Realisation.objects.filter(id_projet=projet)

        for info in infos_specific:
            realisation = realisations.first()  # On suppose qu'il n'y a qu'une seule réalisation
            specific = {
                'nom_projet': info.nom,
                'secteurs': info.id_secteur,
                'objectif_globale': info.objectifs_principals,
                'benef_homme': info.benef_direct_homme,
                'benef_femme': info.benef_direct_femme,
                'total_benef': info.total_benef_direct,
                'partenaire': info.partenaires,
                'depenses': [],
                'ressources_financieres': []
            }

            for depense in depenses:
                specific['depenses'].append({
                    'charge_fonctionnement': depense.salaire_avantages,
                    'equipement': depense.equipement_materiel,
                    'intervention': depense.consommable_divers
                })
            
            if realisation:
                specific['ressources_financieres'].append({
                    'contribution_benef': realisation.contribution_beneficiaire,
                    'apport_partenaire': realisation.contribution_partenaire,
                    'contribution_etat': realisation.part_burkina
                })

            synthese.append(specific)

    context = {
        'synthese': synthese
    }

    # Charger le template
    template_path = 'service/pdf_template_invoice.html'  # Chemin vers ton template
    template = get_template(template_path)
    html = template.render(context)

    # Créer le PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="synthese_projet.pdf"'
    
    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la génération du PDF')

    return response

def generate_word(request):
    projets = Projet.objects.all()
    synthese = []

    for projet in projets:
        infos_specific = InfosSpecific.objects.filter(id_projet=projet)
        depenses = Depense.objects.filter(id_projet=projet)
        realisations = Realisation.objects.filter(id_projet=projet)

        for info in infos_specific:
            specific = {
                'nom_projet': info.nom,
                'secteurs': info.id_secteur,
                'objectif_globale': info.objectifs_principals,
                'benef_homme': info.benef_direct_homme,
                'benef_femme': info.benef_direct_femme,
                'total_benef': info.total_benef_direct,
                'partenaire': info.partenaires,
                'depenses': [],
                'ressources_financieres': []    
            }

            for depense in depenses:
                specific['depenses'].append({
                    'charge_fonctionnement': depense.salaire_avantages,
                    'equipement': depense.equipement_materiel,
                    'intervention': depense.consommable_divers
                })
            
            for realisation in realisations:
                specific['ressources_financieres'].append({
                    'contribution_benef': realisation.contribution_beneficiaire,
                    'apport_partenaire': realisation.contribution_partenaire,
                    'contribution_etat': realisation.part_burkina
                })

            synthese.append(specific)

    # Création du document Word
    doc = Document()
    doc.add_heading("INFORMATIONS GLOBALES", level=1)

    for projet in synthese:
        doc.add_heading(f"Projet N° {synthese.index(projet) + 1}", level=2)
        doc.add_paragraph(f"Intitulé du projet : {projet["nom_projet"]}")
        doc.add_paragraph(f"Secteurs d\'intervention : {projet["secteurs"]}")
        doc.add_paragraph(f"Objectif global : {projet["objectif_globale"]}")
        doc.add_paragraph("Nombre de bénéficiaires :")
        doc.add_paragraph(f"Hommes : {projet["benef_homme"]}")
        doc.add_paragraph(f"Femmes : {projet["benef_femme"]}")
        doc.add_paragraph(f"Total : {projet["total_benef"]}")
        doc.add_paragraph(f"Partenaire financier : {projet["partenaire"]}")

        doc.add_heading("Dépenses du projet :", level=3)
        for depense in projet['depenses']:
            doc.add_paragraph(f"Montant Fonctionnement : {depense["charge_fonctionnement"]}")
            doc.add_paragraph(f"Montant Équipements : {depense["equipement"]}")
            doc.add_paragraph(f"Montant Interventions : {depense["intervention"]}")

        doc.add_heading("Ressources financières :", level=3)
        for ressource in projet['ressources_financieres']:
            doc.add_paragraph(f"Contribution des bénéficiaires : {ressource["contribution_benef"]}")
            doc.add_paragraph(f"Apport des partenaires : {ressource["apport_partenaire"]}")
            doc.add_paragraph(f"Contribution de l\'État : {ressource["contribution_etat"]}")

    # Envoi du fichier Word en réponse
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="synthese_globale.docx"'
    doc.save(response)

    return response

def generate_excel(request):
    projets = Projet.objects.all()
    synthese = []

    for projet in projets:
        infos_specific = InfosSpecific.objects.filter(id_projet=projet)
        depenses = Depense.objects.filter(id_projet=projet)
        realisations = Realisation.objects.filter(id_projet=projet)

        for info in infos_specific:
            specific = {
                'nom_projet': info.nom,
                'secteurs': str(info.id_secteur),
                'objectif_globale': info.objectifs_principals,
                'benef_homme': info.benef_direct_homme,
                'benef_femme': info.benef_direct_femme,
                'total_benef': info.total_benef_direct,
                'partenaire': info.partenaires,
                'changements': info.objectifs_principals,  # Assurez-vous que ce champ existe
                'depenses': [],
                'ressources_financieres': []    
            }

            for depense in depenses:
                specific['depenses'].append({
                    'charge_fonctionnement': depense.salaire_avantages,
                    'equipement': depense.equipement_materiel,
                    'intervention': depense.consommable_divers,
                    'total_depenses': depense.salaire_avantages + depense.equipement_materiel + depense.consommable_divers
                })
            
            for realisation in realisations:
                specific['ressources_financieres'].append({
                    'contribution_benef': realisation.contribution_beneficiaire,
                    'apport_partenaire': realisation.contribution_partenaire,
                    'contribution_etat': realisation.part_burkina,
                    'total_ressources': (realisation.contribution_beneficiaire + 
                                         realisation.contribution_partenaire + 
                                         realisation.part_burkina)
                })

            synthese.append(specific)

    # Création du fichier Excel
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Synthèse des Projets"

    # En-têtes
    headers = [
        "N° D'ordre", "Intitulés des projets", "Liste des secteurs d'intervention couvert par projet",
        "Objectif global du projet", "Nombre de bénéficiaires directes touchés par le projet",
        "Hommes", "Femmes", "Total", "Changements concrets constatés",
        "Montant dépensé pour le Fonctionnement", "Montant dépensé pour les équipements",
        "Montant dépensé pour les interventions", "Montant Total dépensé",
        "Contribution des bénéficiaires et du SED", "Apport des partenaires financiers",
        "Contribution de l'État Burkinabe", "Total des ressources financières du projet",
        "Partenaires Financiers du Projet"
    ]
    worksheet.append(headers)

    # Remplir les données
    for index, projet in enumerate(synthese, start=1):
        total_depenses = sum(dep['total_depenses'] for dep in projet['depenses'])
        total_ressources = sum(res['total_ressources'] for res in projet['ressources_financieres'])

        row = [
            index,
            projet['nom_projet'],
            projet['secteurs'],
            projet['objectif_globale'],
            projet['total_benef'],
            projet['benef_homme'],
            projet['benef_femme'],
            projet['total_benef'],
            projet['changements'],
            total_depenses,  # Total dépenses
            '',  # Montant Équipements (à remplir selon besoin)
            '',  # Montant Interventions (à remplir selon besoin)
            total_depenses,  # Montant Total dépensé
            '',  # Contribution des bénéficiaires et du SED (à remplir selon besoin)
            '',  # Apport des partenaires (à remplir selon besoin)
            '',  # Contribution de l'État (à remplir selon besoin)
            total_ressources,  # Total des ressources financières
            projet['partenaire']
        ]

        worksheet.append(row)

    # Envoi du fichier Excel en réponse
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="synthese_globale.xlsx"'
    workbook.save(response)
    return response


# synthese globale requis
def synthese_globale_requis(request):
    projets = Projet.objects.all()
    synthese = []

    for projet in projets:
        infos_specific = InfosSpecific.objects.filter(id_projet=projet)
        depenses = Depense.objects.filter(id_projet=projet)
        realisations = Realisation.objects.filter(id_projet=projet)

        for info in infos_specific:
            realisation = realisations.first()
            specific = {
                'nom_projet': info.nom,
                'secteurs': info.id_secteur,
                'objectif_globale': info.objectifs_principals,
                'benef_homme': info.benef_direct_homme,
                'benef_femme': info.benef_direct_femme,
                'total_benef': info.total_benef_direct,
                'partenaire': info.partenaires,
                'depenses': [],
                'ressources_financieres': []    
            }

            for depense in depenses:
                specific['depenses'].append({
                    'charge_fonctionnement': depense.salaire_avantages,
                    'equipement': depense.equipement_materiel,
                    'intervention': depense.consommable_divers
                })
            
            for realisation in realisations:
                specific['ressources_financieres'].append({
                    'contribution_benef': realisation.contribution_beneficiaire,
                    'apport_partenaire': realisation.contribution_partenaire,
                    'contribution_etat': realisation.part_burkina
                })

            synthese.append(specific)

    context = {
        'synthese': synthese
    }
    return render(request, 'service/synthese_globale_requis.html', context) 

# generer_synthese
def generer_synthese(request, format):
    # Récupération des données comme dans votre vue originale
    ids = request.GET.get('ids')
    infos_specifiques = []
    
    if ids:
        ids_list = ids.split(',')
        projets = Projet.objects.filter(id__in=ids_list)
        infos = InfosSpecific.objects.filter(id_projet__in=ids_list)

        secteurs = defaultdict(list)
        for info in infos:
            secteur = info.id_secteur
            specific = {
                'secteur': secteur,
                'nom_projet': info.nom,
                'benef_direct_homme': info.benef_direct_homme,
                'benef_direct_femme': info.benef_direct_femme,
                'total_benef': info.total_benef_direct,
                'montant': info.budget,
            }
            secteurs[secteur].append(specific)
        
        for secteur, items in secteurs.items():
            infos_specifiques.append({
                'secteur': secteur,
                'infos': items
            })
    
    if format == 'pdf':
        return generer_pdf(infos_specifiques)
    elif format == 'word':
        return generer_word(infos_specifiques)
    elif format == 'excel':
        return generer_excel(infos_specifiques)

def generer_pdf(infos_specifiques):
    """Génère le PDF de la synthèse"""
    context = {'infos_specifiques': infos_specifiques}
    template_path = 'service/synthese_template.html'
    template = get_template(template_path)
    html = template.render(context)

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="synthese_projets.pdf"'
    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la génération du PDF')
    
    return response

def generer_word(infos_specifiques):
    """Génère le document Word de la synthèse"""
    doc = Document()
    doc.add_heading('Synthèse des Projets', 0)

    # Création du tableau
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    headers = ['Secteur', 'Montant', 'Total Bénéficiaires', 'Bénéficiaires (hommes)', 'Bénéficiaires (femmes)']
    for i, header in enumerate(headers):
        header_cells[i].text = header

    # Données
    for secteur_info in infos_specifiques:
        for info in secteur_info['infos']:
            row_cells = table.add_row().cells
            row_cells[0].text = str(secteur_info['secteur'])
            row_cells[1].text = str(info['montant'])
            row_cells[2].text = str(info['total_benef'])
            row_cells[3].text = str(info['benef_direct_homme'])
            row_cells[4].text = str(info['benef_direct_femme'])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="synthese_projets.docx"'
    doc.save(response)
    return response

def generer_excel(infos_specifiques):
    """Génère le fichier Excel de la synthèse"""
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="synthese_projets.xlsx"'
    
    workbook = xlsxwriter.Workbook(response)
    worksheet = workbook.add_worksheet()

    # Styles
    header_format = workbook.add_format({
        'bold': True,
        'bg_color': '#F0F0F0',
        'border': 1
    })

    # En-têtes
    headers = ['Secteur', 'Montant', 'Total Bénéficiaires', 'Bénéficiaires (hommes)', 'Bénéficiaires (femmes)']
    for col, header in enumerate(headers):
        worksheet.write(0, col, header, header_format)

    # Données
    row = 1
    for secteur_info in infos_specifiques:
        for info in secteur_info['infos']:
            worksheet.write(row, 0, str(secteur_info['secteur']))
            worksheet.write(row, 1, info['montant'])
            worksheet.write(row, 2, info['total_benef'])
            worksheet.write(row, 3, info['benef_direct_homme'])
            worksheet.write(row, 4, info['benef_direct_femme'])
            row += 1

    # Ajuster la largeur des colonnes
    for i, header in enumerate(headers):
        worksheet.set_column(i, i, 15)

    workbook.close()
    return response

# synthese par projet
def synthese_par_projet(request):
    ids = request.GET.get('ids')
    infos_specifiques = []
    # total_montant = 0
    # total_benef = 0
    # total_benef_homme = 0
    # total_benef_femme = 0
    
    if ids:
        ids_list = ids.split(',')
        projets = Projet.objects.filter(id__in=ids_list)
        infos = InfosSpecific.objects.filter(id_projet__in=ids_list)

        # Regrouper les infos par secteur
        secteurs = defaultdict(list)
        for info in infos:
            # total_montant += info.budget
            # total_benef += info.total_benef_direct
            # total_benef_homme += info.benef_direct_homme
            # total_benef_femme += info.benef_direct_femme
            secteur = info.id_secteur
            specific = {
                'secteur': secteur,
                'nom_projet': info.nom,
                'benef_direct_homme': info.benef_direct_homme,
                'benef_direct_femme': info.benef_direct_femme,
                'total_benef': info.total_benef_direct,
                'montant': info.budget,
                # 'total_montant': total_montant,
                # 'total_benef': total_benef,
                # 'total_benef_homme': total_benef_homme,
                # 'total_benef_femme': total_benef_femme
            }
            secteurs[secteur].append(specific)
        
         # Convertir le defaultdict en liste
        for secteur, items in secteurs.items():
            infos_specifiques.append({
                'secteur': secteur,
                'infos': items
            })
       
    else:
        projets = []
        
    context = {
        'projets': projets,
        'infos_specifiques': infos_specifiques
    }
    return render(request, 'service/synthese_par_projet.html', context)

# selection des projets
@login_required(login_url='/login/')
@all_user_required
def selection(request):
    user = request.user
    projects = Projet.objects.all()
    context = {
        'user': user,
        'projects': projects
    }
    return render(request, 'service/selection.html', context)


# recuperer le modele de l'utilisation
User = get_user_model()
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_pdf_situation(request):
    situations = Situation.objects.all()

    audit = {}

    for situation in situations:
        audit_group = Audit.objects.filter(situation=situation)
        audit[situation.id] = [{
            'designation': audits.designation,
            'date_realisation': audits.date_realisation,
            'nom_cabinet': audits.nom_cabinet
        } for audits in audit_group]

    context = {
        'situations': situations,
        'audit': audit
    }
    template_path = 'service/invoice_situation.html'
    template = get_template(template_path)
    html = template.render(context)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="synthese_situation.pdf"'
    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la génération du PDF')
    return response

@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_word_situation(request):
    situations = Situation.objects.all()

    audit = {}

    for situation in situations:
        audit_group = Audit.objects.filter(situation=situation)
        audit[situation.id] = [{
            'designation': audits.designation,
            'date_realisation': audits.date_realisation,
            'nom_cabinet': audits.nom_cabinet
        } for audits in audit_group]

    doc = Document()
    doc.add_heading('SYNTHÈSE GLOBALE SITUATION FISCALE ET SOCIALE', level=1)

    for situation in situations:
        doc.add_heading(f"Utilisateur : { situation.utilisateur } ", level=2)
        doc.add_heading(f"Projet : { situation.id_projet.nom } ", level=2)

        doc.add_heading(f"SITUATIONS FISCALES ET SOCIALES DU SED AU 31 DECEMBRE", level=3)
        doc.add_paragraph(f"Impôts et taxes versées: { situation.impot } FCFA")
        doc.add_paragraph(f"Cotisations Sociales versées : { situation.cotisation } FCFA")
        doc.add_paragraph(f"Autres contributions fiscales versées : { situation.autre_contribution } FCFA")
        doc.add_paragraph(f"Total : { situation.total }")

        doc.add_heading(f"DERNIERS AUDITS COMPTABLES REALISES AU COURS DE L'ANNEE", level=3)
        audits = audit[situation.id]
        if audits:
            for p in audits:
                doc.add_paragraph(f"Désignation projets et programmes audités : {p['designation']} ")
                doc.add_paragraph(f"Date de réalisation : {p['date_realisation']} ")
                doc.add_paragraph(f"Nom du cabinet ayant conduit l'audit comptable : {p['nom_cabinet']} ")
                doc.add_paragraph("----------------------------------------")
        else:
            doc.add_paragraph(f"Désignation projets et programmes audités : {p['designation']} ")
            doc.add_paragraph(f"Date de réalisation : {p['date_realisation']} ")
            doc.add_paragraph(f"Nom du cabinet ayant conduit l'audit comptable : {p['nom_cabinet']} ")
        
        doc.add_paragraph("_________________________________________________________________________________________________________")
    
    output = BytesIO()
    doc.save(output)
    
    output.seek(0)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=synthese_situation.docx'
    response.write(output.getvalue())

    return response

@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_excel_situation(request):
    situations = Situation.objects.all()

    audit = {}

    for situation in situations:
        audit_group = Audit.objects.filter(situation=situation)
        audit[situation.id] = [{
            'designation': audits.designation,
            'date_realisation': audits.date_realisation,
            'nom_cabinet': audits.nom_cabinet
        } for audits in audit_group]
    
    wb = Workbook()
    ws = wb.active

    headers = [
        'Utilisateur', 'Projet', 'Impôts et taxes versées', 'Cotisations Sociales versées', 'Autres contributions fiscales versées', 'Total',
        'Désignation projets et programmes audités', 'Date de réalisation', 'Nom du cabinet ayant conduit l\'audit comptable'
    ]
    ws.append(headers)

    for situation in situations:
        #audits = audit[situation.id]
        audits = audit.get(situation.id, [])

        if audits:
            designation = safe_join([audit['designation'] for audit in audits])
            date_realisation = safe_join([audit['date_realisation'] for audit in audits])
            nom_cabinet = safe_join([audit['nom_cabinet'] for audit in audits])
        else:
            designation = ''
            date_realisation = ''
            nom_cabinet = ''

        row = [
            situation.utilisateur, situation.id_projet.nom, situation.impot, situation.cotisation, situation.autre_contribution, situation.total,
            designation, date_realisation, nom_cabinet
        ]
        ws.append(row)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=synthese_situation.xlsx'

    wb.save(response)

    return response

# Generer le pdf de generale
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_pdf_generale(request):
    infos_test = InfosGenerale.objects.all()

    partenariats_dict = {}
    objectifs_dict = {}

    for info in infos_test:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_dict[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_dict[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]

    context = {
        'infos_test': infos_test,
        'partenariats_dict': partenariats_dict,
        'objectifs_dict': objectifs_dict
    }
    template_path = 'service/invoice_general.html'
    template = get_template(template_path)
    html = template.render(context)
    # Créer un objet HttpResponse avec le type de contenu PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="synthese_globale.pdf"'
    # Convertir le template HTML en PDF
    pisa_status = pisa.CreatePDF(html, dest=response)
    # Si la conversion a réussi, retourner la réponse avec le PDF généré
    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la génération du PDF')
    return response

# Generer le word de generale
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_word_generale(request):
    infos_test = InfosGenerale.objects.all()

    partenariats_dict = {}
    objectifs_dict = {}

    for info in infos_test:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_dict[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_dict[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]

    doc = Document()
    doc.add_heading('SYNTHÈSE GLOBALE INFORMATIONS GÉNÉRALES', level=1)

    for info in infos_test:
        doc.add_heading(f"Créateur du projet : {info.utilisateur.username}", level=2)

        # infos organisation
        doc.add_heading(f"Informations sur l'organisation", level=3)
        doc.add_paragraph(f"Nom de l'organisation : {info.nom_org}")
        doc.add_paragraph(f"Nature de l'organisation : {info.nature_org}")
        doc.add_paragraph(f"Sigle : {info.sigle}")
        doc.add_paragraph(f"Pays d'origine  : {info.pays_origine}")
        #
        doc.add_heading(f"Adresses du siege de l'organisation", level=3)
        doc.add_paragraph(f"Région : {info.region}")
        doc.add_paragraph(f"Province : {info.province}")
        doc.add_paragraph(f"Commune : {info.commune}")
        doc.add_paragraph(f"Ville/Secteur : {info.village}")
        doc.add_paragraph(f"Boite postale : {info.boite_postale}")
        doc.add_paragraph(f"Numéro de téléphone fixe : {info.numb_fixe}")
        doc.add_paragraph(f"Numéro de téléphone mobile : {info.numb_mobile}")
        doc.add_paragraph(f"Adresse mail professionnelle : {info.adresse_mail}")
        doc.add_paragraph(f"Site web : {info.site_web}")
        #
        doc.add_heading(f"Responsable de l'organisation", level=3)
        doc.add_paragraph(f"Nom et Prénom(s) : {info.nom_complet_resp}")
        doc.add_paragraph(f"Nationalité : {info.nationalite_resp}")
        doc.add_paragraph(f"Fonction(Président,...) : {info.fonction_resp}")
        doc.add_paragraph(f"Numéro fixe : {info.numb_fixe_resp}")
        doc.add_paragraph(f"Numéro mobile : {info.numb_mobile_resp}")
        #
        doc.add_heading(f"Gouvernance interne de l'association : Tenue des rencontres statuaires des instances de l'organisation", level=3)
        doc.add_paragraph(f"Dernier renouvèlement des Instances dirigeantes : {info.renou_instance}")
        doc.add_paragraph(f"Dernière Assemblée Générale Ordinaire  : {info.assem_general}")
        doc.add_paragraph(f"Dernière session statutaire du bureau exécutif : {info.session_statut}")
        doc.add_paragraph(f"Durée du mandat du bureau exécutif : {info.mandat_bureau} ans")
        #
        doc.add_heading(f"Objectifs principaux de l'organisation", level=3)
        objectifs = objectifs_dict[info.id]
        objectifs_str = '\n'.join([o['objectifs'] for o in objectifs]) + '\n'
        doc.add_paragraph(f"Objectifs : \n {objectifs_str}")
        #
        doc.add_heading(f"Groupes cibles specifique", level=3)
        doc.add_paragraph(f"Groupes : {info.groupes_cibles}")
        #
        doc.add_heading(f"Personnel employe", level=3)
        doc.add_heading(f"Nombre total du personnel", level=4)
        doc.add_paragraph(f"Hommes : {info.total_pers_homme}")
        doc.add_paragraph(f"Femmes : {info.total_pers_femme}")
        doc.add_heading(f"Employés nationaux Contrat à Durée Indéterminée (CDI)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_nation_cdi_homme}")
        doc.add_paragraph(f"Femmes : {info.em_nation_cdi_femme}")
        doc.add_heading(f"Employés nationaux Contrat à Durée déterminée (CDD)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_nation_cdd_homme}")
        doc.add_paragraph(f"Femmes : {info.em_nation_cdd_femme}")
        doc.add_heading(f"Employés expatriés Contrat à Durée Indéterminée (CDI)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_expa_cdi_homme}")
        doc.add_paragraph(f"Femmes : {info.em_expa_cdi_femme}")
        doc.add_heading(f"Employés expatriés Contrat à Durée déterminée (CDD)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_expa_cdd_homme}")
        doc.add_paragraph(f"Femmes : {info.em_expa_cdd_femme}")
        #
        doc.add_heading(f"Bénévoles ou volontaires", level=3)
        doc.add_heading(f"Bénévoles ou volontaires Nationaux", level=4)
        doc.add_paragraph(f"Hommes : {info.benevol_nation_homme}")
        doc.add_paragraph(f"Femmes : {info.benevol_nation_femme}")
        doc.add_heading(f"Bénévoles ou volontaires Expatriés", level=4)
        doc.add_paragraph(f"Hommes : {info.benevol_expa_homme}")
        doc.add_paragraph(f"Femmes : {info.benevol_expa_femme}")
        #
        doc.add_heading(f"Personnel de l'Administration publique en détachement", level=3)
        doc.add_paragraph(f"Hommes : {info.personnel_admin_homme}")
        doc.add_paragraph(f"Femmes : {info.personnel_admin_femme}")
        #
        doc.add_heading(f"Partenariats / collaborations", level=3)
        partenariats = partenariats_dict[info.id]
        if partenariats:
            for p in partenariats:
                doc.add_paragraph(f"Nom du partenaire : {p['nom']}")
                doc.add_paragraph(f"N° de convention du partenariat / protocole d'entente : {p['numero']}")
                doc.add_paragraph(f"Date de début d'effet : {p['date_debut']}")
                doc.add_paragraph(f"Date de fin d'effet : {p['date_fin']}")
                doc.add_paragraph("----------------------------------------")
        else:
            doc.add_paragraph(f"Nom du partenaire : ")
            doc.add_paragraph(f"N° de convention de partenariat / protocole d'entente : ")
            doc.add_paragraph(f"Date de début d'effet : ")
            doc.add_paragraph(f"Date de fin d'effet : ")

        doc.add_paragraph("_________________________________________________________________________________________________________")
    output = BytesIO()
    doc.save(output)

    output.seek(0)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=synthese_generales.docx'
    response.write(output.getvalue())

    return response


# Generer le excel de generale
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_excel_generale(request):
    infos_test = InfosGenerale.objects.all()

    partenariats_dict = {}
    objectifs_dict = {}

    for info in infos_test:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_dict[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_dict[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]

    wb = Workbook()
    ws = wb.active

    headers = [
        'Créateur du projet', 'Nom de l\'organisation', 'Nature de l\'organisation', 'Sigle', 'Pays d\'origine', 'Région', 'Province', 'Commune', 'Ville/Secteur',
        'Boite postale', 'Numéro de téléphone fixe', 'Numéro de téléphone mobile', 'Adresse mail professionnelle', 'Site web', 'Nom et Prénom(s)', 'Nationalité',
        'Fonction(Président,...)', 'Numéro fixe', 'Numéro mobile', 'Dernier renouvèlement des Instances dirigeantes', 'Dernière Assemblée Générale Ordinaire',
        'Dernière session statutaire du bureau exécutif', 'Durée du mandat du bureau exécutif', 'Objectifs', 'Groupes cibles specifique', 'Hommes Nombre total du personnel', 'Femmes Nombre total du personnel', 'Hommes Employés nationaux Contrat à Durée Indéterminée (CDI)',
        'Femmes Employés nationaux Contrat à Durée Indéterminée (CDI)', 'Hommes Employés nationaux Contrat à Durée déterminée (CDD)', 'Femmes Employés nationaux Contrat à Durée déterminée (CDD)',
        'Hommes Employés expatriés Contrat à Durée Indéterminée (CDI)', 'Femmes Employés expatriés Contrat à Durée Indéterminée (CDI)', 'Hommes Employés expatriés Contrat à Durée déterminée (CDD)',
        'Femmes Employés expatriés Contrat à Durée déterminée (CDD)', 'Hommes Bénévoles ou volontaires Nationaux', 'Femmes Bénévoles ou volontaires Nationaux',
        'Hommes Bénévoles ou volontaires Expatriés', 'Femmes Bénévoles ou volontaires Expatriés', 'Hommes Personnel de l\'Administration publique en détachement', 'Femmes Personnel de l\'Administration publique en détachement',
        'Nom partenariats', 'N° de convention de partenariat / protocole d\'entente', 'Date de début d\'effet', 'Date de fin d\'effet'
    ]
    ws.append(headers)
    

    # remplir les donnees
    for info in infos_test:
        partenariats = partenariats_dict[info.id]
        objectifs = objectifs_dict[info.id]

        objectifs_str = safe_join([o['objectifs'] for o in objectifs])

        if partenariats:
            for p in partenariats:
                row = [
                    info.utilisateur.username, info.nom_org , info.nature_org, info.sigle, info.pays_origine, info.region, info.province, info.commune,
                    info.village, info.boite_postale, info.numb_fixe, info.numb_mobile, info.adresse_mail, info.site_web, info.nom_complet_resp,
                    info.nationalite_resp, info.fonction_resp, info.numb_fixe_resp, info.numb_mobile_resp, info.renou_instance, info.assem_general,
                    info.session_statut, info.mandat_bureau,  objectifs_str, info.groupes_cibles, info.total_pers_homme, info.total_pers_femme, info.em_nation_cdi_homme, info.em_nation_cdi_femme,
                    info.em_nation_cdd_homme, info.em_nation_cdd_femme, info.em_expa_cdi_homme, info.em_expa_cdi_femme, info.em_expa_cdd_homme,
                    info.em_expa_cdd_femme, info.benevol_nation_homme, info.benevol_nation_femme, info.benevol_expa_homme, info.benevol_expa_femme,
                    info.personnel_admin_homme, info.personnel_admin_femme, p['nom'], p['numero'], p['date_debut'], p['date_fin']
                ]
                ws.append(row)
        else:
            row = [
                info.utilisateur.username, info.nom_org, info.nature_org, info.sigle, info.pays_origine, info.region, info.province, info.commune,
                    info.village, info.boite_postale, info.numb_fixe, info.numb_mobile, info.adresse_mail, info.site_web, info.nom_complet_resp,
                    info.nationalite_resp, info.fonction_resp, info.numb_fixe_resp, info.numb_mobile_resp, info.renou_instance, info.assem_general,
                    info.session_statut, info.mandat_bureau,  objectifs_str, info.groupes_cibles, info.total_pers_homme, info.total_pers_femme, info.em_nation_cdi_homme, info.em_nation_cdi_femme,
                    info.em_nation_cdd_homme, info.em_nation_cdd_femme, info.em_expa_cdi_homme, info.em_expa_cdi_femme, info.em_expa_cdd_homme,
                    info.em_expa_cdd_femme, info.benevol_nation_homme, info.benevol_nation_femme, info.benevol_expa_homme, info.benevol_expa_femme,
                    info.personnel_admin_homme, info.personnel_admin_femme, '', '', '', ''
            ]
            ws.append(row)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=synthese_generales.xlsx'

    wb.save(response)

    return response


# Generer le pdf de synthese de specifique
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_pdf_specifique(request):
    specific_test = InfosSpecific.objects.all()
    objectif_dict = {}
    resultat_dict = {}

    for specific in specific_test:
        objectif_group = Objectif.objects.filter(id_infos_specifique=specific)
        objectif_dict[specific.id] = [{
            'objectifs': objectifs.objectifs
        } for objectifs in objectif_group ]

        resultat_group = Resultat.objects.filter(id_specific=specific)
        resultat_dict[specific.id] = [{
            'resultats': resultat.resultats
        } for resultat in resultat_group]
    
    context = {
        'specific_test': specific_test,
        'objectif_dict': objectif_dict,
        'resultat_dict': resultat_dict
    }

    template_path = 'service/invoice_specifique.html'
    template = get_template(template_path)
    html = template.render(context)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachement; filename="synthese_specifique.pdf"'
    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la generation du PDF')
    return response

# Generer word information specifique
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_word_specifique(request):
    specific_test = InfosSpecific.objects.all()
    objectif_dict = {}
    resultat_dict = {}

    for specific in specific_test:
        objectif_group = Objectif.objects.filter(id_infos_specifique=specific)
        objectif_dict[specific.id] = [{
            'objectifs': objectifs.objectifs
        } for objectifs in objectif_group ]

        resultat_group = Resultat.objects.filter(id_specific=specific)
        resultat_dict[specific.id] = [{
            'resultats': resultat.resultats
        } for resultat in resultat_group]
    
    doc = Document()
    doc.add_heading("SYNTHESE GLOBALE INFORMATIONS SPECIFIQUES", level=1)

    for specific in specific_test:
        doc.add_heading(f"Nom du projet : { specific.nom }", level=2)

        doc.add_paragraph(f"Date de debut du projet : { specific.date_debut }")
        doc.add_paragraph(f"Date de fin du projet : : { specific.date_fin }")
        doc.add_paragraph(f"Cout : { specific.budget }")
        doc.add_paragraph(f"Nombre de beneficiaires direct hommes : { specific.benef_direct_homme }")
        doc.add_paragraph(f"Nombre de beneficiaires direct femmes : { specific.benef_direct_femme }")
        doc.add_paragraph(f"Objectif principale : { specific.objectifs_principals }")
        doc.add_heading(f"Objectifs secondaires", level=3)

        objectifs = objectif_dict[specific.id]
        objectifs_str = '\n'.join([o['objectifs'] for o in objectifs]) + '\n'
        doc.add_paragraph(f" - \n { objectifs_str }")

        doc.add_heading("Resultats", level=3)
        
        resultats = resultat_dict[specific.id]
        resultats_str = '\n'.join([o['resultats'] for o in resultats]) + '\n'
        doc.add_paragraph(f" -  \n { resultats_str } ")

        doc.add_paragraph(f"Partenaire financier principale : { specific.partenaires }")
        doc.add_paragraph("_________________________________________________________________________________________________________")

    output = BytesIO()
    doc.save(output)

    output.seek(0)

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response['Content-Disposition'] = 'attachement; filename="synthese_specifique.docx"'
    response.write(output.getvalue())

    return response

@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_excel_specifique(request):
    specific_test = InfosSpecific.objects.all()

    objectif_dict = {}
    resultat_dict = {}

    for specific in specific_test:
        objectif_group = Objectif.objects.filter(id_infos_specifique=specific)
        objectif_dict[specific.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group ]

        resultat_group = Resultat.objects.filter(id_specific=specific)
        resultat_dict[specific.id] = [{
            'resultats': resultat.resultats
        } for resultat in resultat_group]

    wb = Workbook()
    ws = wb.active

    headers = [
        'Nom du projet', 'Date de debut du projet', 'Date de fin projet', 'Cout', 'Nombre de beneficiaire directs hommes',
        'Nombre de beneficiaire direct femmes', 'Objectif principale', 'Objectifs secondaires', 'Resultats', 'Partenaire financier principal'
    ]
    ws.append(headers)

    for specific in specific_test:
        objectifs = objectif_dict[specific.id]
        resultats = resultat_dict[specific.id]

        objectifs_str = ', '.join([o['objectifs'] for o in objectifs]) + '\n'
        resultats_str = ', '.join([o['resultats'] for o in resultats]) + '\n'

        row = [
            specific.nom, specific.date_debut, specific.date_fin, specific.cout, specific.benef_direct_homme,
            specific.benef_direct_femme, specific.objectifs_principals, objectifs_str, resultats_str, specific.partenaires
        ]
        ws.append(row)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachement; filename=synthese_specifique.xlsx'

    wb.save(response)
    return response

#
@login_required(login_url='/login/')
@responsable_charger_projet_required
def globale_globale(request):
    # ------------------------------------- Information generales ----------------------
    infos_test = InfosGenerale.objects.all()

    partenariats_dict = {}
    objectifs_dict = {}

    for info in infos_test:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_dict[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_dict[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]

    
    # ----------------------------------- Informations specifique --------------------
    specifiques = InfosSpecific.objects.all()
    objectifs_specific = {}
    resultats_specific = {}

    for specifique in specifiques:
        objectifs = Objectif.objects.filter(id_infos_specifique=specifique)
        objectifs_specific[specifique.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectifs ]

        resultats = Resultat.objects.filter(id_specific=specifique)
        resultats_specific[specifique.id] = [{
            'resultats': resultat.resultats
        } for resultat in resultats ]
    
    # -------------------- Planification ------------------------
    distinct_titres = Activite.objects.values_list('titre', flat=True).distinct()
    activites_groupees = []
    for titre in distinct_titres:
        activite_groupe = Activite.objects.filter(titre=titre)

        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        region = safe_join(activite_groupe.values_list('region', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))

        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))

        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        partenaireP = safe_join(activite_groupe.values_list('partenaires', flat=True))
             
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_activite=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'partenaireP': partenaireP,
            'partenaires': list(partenaires_groupes.values())
        })

# --------------------------------- Realisation ------------------------
    distinct_titres_realisation = Realisation.objects.values_list('titre', flat=True).distinct()
    activites_groupees_realisation = []
    for titre_realisation in distinct_titres_realisation:
        activite_groupe_realisation = Realisation.objects.filter(titre=titre_realisation)
        
        commune_realisation = safe_join(activite_groupe_realisation.values_list('commune', flat=True))
        province_realisation = safe_join(activite_groupe_realisation.values_list('province', flat=True))
        region_realisation = safe_join(activite_groupe_realisation.values_list('region', flat=True))
        paroisse_realisation = safe_join(activite_groupe_realisation.values_list('paroisse', flat=True))
        unite_physique_realisation = safe_join(activite_groupe_realisation.values_list('unite_physique', flat=True))
        quantite_prevue_realisation = safe_sum(activite_groupe_realisation.values_list('quantite_prevue', flat=True))
        periode_prevue_debut_realisation = safe_join(activite_groupe_realisation.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin_realisation = safe_join(activite_groupe_realisation.values_list('periode_prevue_fin', flat=True))
        responsable_realisation = safe_join(activite_groupe_realisation.values_list('responsable', flat=True))
        cout_realisation_realisation = safe_sum(activite_groupe_realisation.values_list('cout_realisation', flat=True))
        contribution_beneficiaire_realisation = safe_sum(activite_groupe_realisation.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire_realisation = safe_sum(activite_groupe_realisation.values_list('contribution_partenaire', flat=True))
        total_benef_direct_realisation = safe_sum(activite_groupe_realisation.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme_realisation = safe_sum(activite_groupe_realisation.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme_realisation = safe_sum(activite_groupe_realisation.values_list('nbre_benef_direct_femme', flat=True))
        partenaireP_realisation = safe_join(activite_groupe_realisation.values_list('partenaires', flat=True))
        
        
        partenaires_realisation = {}
        for activite in activite_groupe_realisation:
            partenaires_activite = Partenaire.objects.filter(id_realisation=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes_realisation = {}
        for partenaire in partenaires_realisation:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees_realisation.append({
            'commune': commune_realisation,
            'province': province_realisation,
            'region': region_realisation,
            'paroisse': paroisse_realisation,
            'titre': titre_realisation,
            'unite_physique': unite_physique_realisation,
            'quantite_prevue': quantite_prevue_realisation,
            'periode_prevue_debut': periode_prevue_debut_realisation,
            'periode_prevue_fin': periode_prevue_fin_realisation,
            'responsable': responsable_realisation,
            'cout_realisation': cout_realisation_realisation,
            'contribution_beneficiaire': contribution_beneficiaire_realisation,
            'contribution_partenaire': contribution_partenaire_realisation,
            'total_benef_direct': total_benef_direct_realisation,
            'nbre_benef_direct_homme': nbre_benef_direct_homme_realisation,
            'nbre_benef_direct_femme': nbre_benef_direct_femme_realisation,
            'partenaireP': partenaireP_realisation,
            'partenaires': list(partenaires_groupes_realisation.values())
        })

    # ----------------------------------- Depenses -------------
    depenses = Depense.objects.all()

    # --------------------------------- Situation -----------------------
    situations = Situation.objects.all()
    audits = {}

    for situation in situations:
        audit_group = Audit.objects.filter(situation=situation)
        audits[situation.id] = [{
            'designation': audits.designation,
            'date_realisation': audits.date_realisation,
            'nom_cabinet': audits.nom_cabinet
        } for audits in audit_group]

    context = {
        'infos_generale': infos_test,
        'partenariat_generales': partenariats_dict,
        'objectifs_generales': objectifs_dict,
        'infos_specific': specifiques,
        'objectifs_specific': objectifs_specific,
        'resultats_specific': resultats_specific,
        'activites_groupees_planification': activites_groupees,
        'activites_groupees_realisation': activites_groupees_realisation,
        'depenses': depenses,
        'situations': situations,
        'audits': audits
    }
    return render(request, 'service/globale_globale.html', context)
#
@login_required(login_url='/login/')
@responsable_charger_projet_required
def globale_specifique(request):
    specifiques = InfosSpecific.objects.all()
    objectifs_dict = {}
    resultats_dict = {}

    for specifique in specifiques:
        objectifs = Objectif.objects.filter(id_infos_specifique=specifique)
        objectifs_dict[specifique.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectifs ]

        resultats = Resultat.objects.filter(id_specific=specifique)
        resultats_dict[specifique.id] = [{
            'resultats': resultat.resultats
        } for resultat in resultats ]

    context = {
        'specifiques': specifiques,
        'objectifs_dict': objectifs_dict,
        'resultats_dict': resultats_dict
    }
    return render(request, 'service/globale_specifique.html', context)

#
@login_required(login_url='/login/')
@responsable_charger_projet_required
def globale_depense(request):
    depenses = Depense.objects.all()
    depense_dict = {}

    for depense in depenses:
        autre = AutreDepense.objects.filter(depense=depense)
        depense_dict[depense.id] = [{
            'intitule': depens.intitule,
            'prix': depens.prix
        } for depens in autre]

    context = {
        'depenses': depenses,
        'depense_dict': depense_dict
    }
    return render(request, 'service/globale_depense.html', context)

# generer un PDF pour la synthese globale
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_pdf_globale(request):
    # Information generales
    infos_generale = InfosGenerale.objects.all()

    partenariats_generale = {}
    objectifs_generale = {}

    for info in infos_generale:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_generale[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_generale[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]
    # Information specifique
    infos_specific = InfosSpecific.objects.all()
    objectif_specific = {}
    resultat_specific = {}

    for specific in infos_specific:
        objectif_group = Objectif.objects.filter(id_infos_specifique=specific)
        objectif_specific[specific.id] = [{
            'objectifs': objectifs.objectifs
        } for objectifs in objectif_group ]

        resultat_group = Resultat.objects.filter(id_specific=specific)
        resultat_specific[specific.id] = [{
            'resultats': resultat.resultats
        } for resultat in resultat_group]
    # Planification
    distinct_titres_planification = Activite.objects.values_list('titre', flat=True).distinct()
    activites_groupees_planification = []
    for titre in distinct_titres_planification:
        activite_groupe = Activite.objects.filter(titre=titre)
        
        region = safe_join(activite_groupe.values_list('region', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        partenaires = safe_join(activite_groupe.values_list('partenaires', flat=True))
        
        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)

        activites_groupees_planification.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre':titre,
            'region': region,
            'province': province,
            'commune': commune,
            'paroisse': paroisse,
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'total_benef_direct': total_benef_direct,
            'partenaires': partenaires
        })
    # Realisation
    distinct_titres_realisation = Realisation.objects.values_list('titre', flat=True).distinct()
    activites_groupees_realisation = []
    for titre in distinct_titres_realisation:
        activite_groupe = Realisation.objects.filter(titre=titre)
    
        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        region = safe_join(activite_groupe.values_list('region', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = safe_join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = safe_join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = safe_join(activite_groupe.values_list('responsable', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        partenaireP = safe_join(activite_groupe.values_list('partenaires', flat=True))

        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)

        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_realisation=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees_realisation.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre': titre,
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'partenaireP': partenaireP,
            'partenaires': list(partenaires_groupes.values())
        })
    # Depense
    depenses = Depense.objects.all()
    # Situation
    situations = Situation.objects.all()

    audits = {}

    for situation in situations:
        audit_group = Audit.objects.filter(situation=situation)
        audits[situation.id] = [{
            'designation': audits.designation,
            'date_realisation': audits.date_realisation,
            'nom_cabinet': audits.nom_cabinet
        } for audits in audit_group]


    context = {
        'infos_generale': infos_generale,
        'partenariats_generale': partenariats_generale,
        'objectifs_generale': objectifs_generale,
        'infos_specific': infos_specific,
        'objectif_specific': objectif_specific,
        'resultat_specific': resultat_specific,
        'activites_groupees_planification': activites_groupees_planification,
        'activites_groupees_realisation': activites_groupees_realisation,
        'depenses': depenses,
        'situations': situations,
        'audits': audits
    }

    template_path = 'service/invoice_globale.html'
    template = get_template(template_path)
    html = template.render(context)
    # Créer un objet HttpResponse avec le type de contenu PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="globale_globale.pdf"'
    # Convertir le template HTML en PDF
    pisa_status = pisa.CreatePDF(html, dest=response)
    # Si la conversion a réussi, retourner la réponse avec le PDF généré
    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la génération du PDF')
    return response

# generer un word pour la synthese globale
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_word_globale(request):
    # Information generale
    infos_generale = InfosGenerale.objects.all()

    partenariats_generale = {}
    objectifs_generale = {}

    for info in infos_generale:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_generale[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_generale[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]

    # Information specifique
    infos_specific = InfosSpecific.objects.all()
    objectif_specific = {}
    resultat_specific = {}

    for specific in infos_specific:
        objectif_group = Objectif.objects.filter(id_infos_specifique=specific)
        objectif_specific[specific.id] = [{
            'objectifs': objectifs.objectifs
        } for objectifs in objectif_group ]

        resultat_group = Resultat.objects.filter(id_specific=specific)
        resultat_specific[specific.id] = [{
            'resultats': resultat.resultats
        } for resultat in resultat_group]

    # Planification
    activites_groupees_planification = []
    distinct_titres_planification = Activite.objects.values_list('titre', flat=True).distinct()

    for titre in distinct_titres_planification:
        activite_groupe = Activite.objects.filter(titre=titre)

        region = safe_join(activite_groupe.values_list('region', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        partenaires = safe_join(activite_groupe.values_list('partenaires', flat=True))

        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)
        
        activites_groupees_planification.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre': titre,
            'region': region,
            'province': province,
            'commune': commune,
            'paroisse': paroisse,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'total_benef_direct': total_benef_direct,
            'partenaires': partenaires          
        })
    # Realisation
    distinct_titres_realisation = Realisation.objects.values_list('titre', flat=True).distinct()
    ##################
    activites_groupees_realisation = []
    for titre in distinct_titres_realisation:
        activite_groupe = Realisation.objects.filter(titre=titre)

        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        region = safe_join(activite_groupe.values_list('region', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = safe_join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = safe_join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = safe_join(activite_groupe.values_list('responsable', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        partenaireP = safe_join(activite_groupe.values_list('partenaires', flat=True))
        
        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)

        activites_groupees_realisation.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre': titre,
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'partenaireP': partenaireP
            })
    # Situation
    situations = Situation.objects.all()

    audits = {}

    for situation in situations:
        audit_group = Audit.objects.filter(situation=situation)
        audits[situation.id] = [{
            'designation': audits.designation,
            'date_realisation': audits.date_realisation,
            'nom_cabinet': audits.nom_cabinet
        } for audits in audit_group]

    # Depense
    depenses = Depense.objects.all()

    # Ajoute des informations au fichier word
    doc = Document()
    # Informations generales
    doc.add_heading('SYNTHESE INFORMATIONS GENERALES', level=1)

    for info in infos_generale:
        doc.add_heading(f"Créateur du projet : {info.utilisateur.username}", level=2)

        # infos organisation
        doc.add_heading(f"Informations sur l'organisation", level=3)
        doc.add_paragraph(f"Nom de l'organisation : {info.nom_org}")
        doc.add_paragraph(f"Nature de l'organisation : {info.nature_org}")
        doc.add_paragraph(f"Sigle : {info.sigle}")
        doc.add_paragraph(f"Pays d'origine  : {info.pays_origine}")
        #
        doc.add_heading(f"Adresses du siege de l'organisation", level=3)
        doc.add_paragraph(f"Région : {info.region}")
        doc.add_paragraph(f"Province : {info.province}")
        doc.add_paragraph(f"Commune : {info.commune}")
        doc.add_paragraph(f"Ville/Secteur : {info.village}")
        doc.add_paragraph(f"Boite postale : {info.boite_postale}")
        doc.add_paragraph(f"Numéro de téléphone fixe : {info.numb_fixe}")
        doc.add_paragraph(f"Numéro de téléphone mobile : {info.numb_mobile}")
        doc.add_paragraph(f"Adresse mail professionnelle : {info.adresse_mail}")
        doc.add_paragraph(f"Site web : {info.site_web}")
        #
        doc.add_heading(f"Responsable de l'organisation", level=3)
        doc.add_paragraph(f"Nom et Prénom(s) : {info.nom_complet_resp}")
        doc.add_paragraph(f"Nationalité : {info.nationalite_resp}")
        doc.add_paragraph(f"Fonction(Président,...) : {info.fonction_resp}")
        doc.add_paragraph(f"Numéro fixe : {info.numb_fixe_resp}")
        doc.add_paragraph(f"Numéro mobile : {info.numb_mobile_resp}")
        #
        doc.add_heading(f"Gouvernance interne de l'association : Tenue des rencontres statuaires des instances de l'organisation", level=3)
        doc.add_paragraph(f"Dernier renouvèlement des Instances dirigeantes : {info.renou_instance}")
        doc.add_paragraph(f"Dernière Assemblée Générale Ordinaire  : {info.assem_general}")
        doc.add_paragraph(f"Dernière session statutaire du bureau exécutif : {info.session_statut}")
        doc.add_paragraph(f"Durée du mandat du bureau exécutif : {info.mandat_bureau} ans")
        #
        doc.add_heading(f"Objectifs principaux de l'organisation", level=3)
        objectifs = objectifs_generale[info.id]
        objectifs_str = '\n'.join([o['objectifs'] for o in objectifs]) + '\n'
        doc.add_paragraph(f"\n {objectifs_str}")
        #
        doc.add_heading(f"Groupes cibles specifique", level=3)
        doc.add_paragraph(f"Groupes : {info.groupes_cibles}")
        #
        doc.add_heading(f"Personnel employe", level=3)
        doc.add_heading(f"Nombre total du personnel", level=4)
        doc.add_paragraph(f"Hommes : {info.total_pers_homme}")
        doc.add_paragraph(f"Femmes : {info.total_pers_femme}")
        doc.add_heading(f"Employés nationaux Contrat à Durée Indéterminée (CDI)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_nation_cdi_homme}")
        doc.add_paragraph(f"Femmes : {info.em_nation_cdi_femme}")
        doc.add_heading(f"Employés nationaux Contrat à Durée déterminée (CDD)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_nation_cdd_homme}")
        doc.add_paragraph(f"Femmes : {info.em_nation_cdd_femme}")
        doc.add_heading(f"Employés expatriés Contrat à Durée Indéterminée (CDI)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_expa_cdi_homme}")
        doc.add_paragraph(f"Femmes : {info.em_expa_cdi_femme}")
        doc.add_heading(f"Employés expatriés Contrat à Durée déterminée (CDD)", level=4)
        doc.add_paragraph(f"Hommes : {info.em_expa_cdd_homme}")
        doc.add_paragraph(f"Femmes : {info.em_expa_cdd_femme}")
        #
        doc.add_heading(f"Bénévoles ou volontaires", level=3)
        doc.add_heading(f"Bénévoles ou volontaires Nationaux", level=4)
        doc.add_paragraph(f"Hommes : {info.benevol_nation_homme}")
        doc.add_paragraph(f"Femmes : {info.benevol_nation_femme}")
        doc.add_heading(f"Bénévoles ou volontaires Expatriés", level=4)
        doc.add_paragraph(f"Hommes : {info.benevol_expa_homme}")
        doc.add_paragraph(f"Femmes : {info.benevol_expa_femme}")
        #
        doc.add_heading(f"Personnel de l'Administration publique en détachement", level=3)
        doc.add_paragraph(f"Hommes : {info.personnel_admin_homme}")
        doc.add_paragraph(f"Femmes : {info.personnel_admin_femme}")
        #
        doc.add_heading(f"Partenariats / collaborations", level=3)
        partenariats = partenariats_generale[info.id]
        if partenariats:
            for p in partenariats:
                doc.add_paragraph(f"Nom du partenaire : {p['nom']}")
                doc.add_paragraph(f"N° de convention du partenariat / protocole d'entente : {p['numero']}")
                doc.add_paragraph(f"Date de début d'effet : {p['date_debut']}")
                doc.add_paragraph(f"Date de fin d'effet : {p['date_fin']}")
                doc.add_paragraph("----------------------------------------")
        else:
            doc.add_paragraph(f"Nom du partenaire : ")
            doc.add_paragraph(f"N° de convention de partenariat / protocole d'entente : ")
            doc.add_paragraph(f"Date de début d'effet : ")
            doc.add_paragraph(f"Date de fin d'effet : ")

        doc.add_paragraph("_________________________________________________________________________________________________________")
    
    doc.add_paragraph()
    doc.add_heading("SYNTHESE INFORMATIONS SPECIFIQUES", level=1)

    for specific in infos_specific:
        doc.add_heading(f"Nom du projet : { specific.nom }", level=2)

        doc.add_paragraph(f"Date de debut du projet : { specific.date_debut }")
        doc.add_paragraph(f"Date de fin du projet : : { specific.date_fin }")
        doc.add_paragraph(f"Cout : { specific.budget }")
        doc.add_paragraph(f"Nombre de beneficiaires direct hommes : { specific.benef_direct_homme }")
        doc.add_paragraph(f"Nombre de beneficiaires direct femmes : { specific.benef_direct_femme }")
        doc.add_paragraph(f"Objectif principale : { specific.objectifs_principals }")
        doc.add_heading(f"Objectifs secondaires", level=3)

        objectifs = objectif_specific[specific.id]
        objectifs_str = '\n'.join([o['objectifs'] for o in objectifs]) + '\n'
        doc.add_paragraph(f"\n - { objectifs_str }")

        doc.add_heading("Resultats", level=3)
        
        resultats = resultat_specific[specific.id]
        resultats_str = '\n'.join([o['resultats'] for o in resultats]) + '\n'
        doc.add_paragraph(f"\n - { resultats_str } ")

        doc.add_paragraph(f"Partenaire financier principale : { specific.partenaires }")
        doc.add_paragraph("_________________________________________________________________________________________________________")
    
    doc.add_paragraph()
    doc.add_heading("SYNTHESE PLANIFICATION", level=1)

    for activite in activites_groupees_planification:
        doc.add_heading(f"{activite['projet']}")
        
        doc.add_heading(f"Activitée", level=2)
        doc.add_paragraph(f"{activite['titre']}")

        doc.add_heading(f"Secteur", level=2)
        doc.add_paragraph(f"{activite['secteur']}")

        doc.add_heading(f"Domaine", level=2)
        doc.add_paragraph(f"{activite['domaine']}")

        doc.add_heading(f"Région", level=2)
        doc.add_paragraph(f"{activite['region']}")

        doc.add_heading(f"Province", level=2)
        doc.add_paragraph(f"{activite['province']}")

        doc.add_heading(f"Commune", level=2)
        doc.add_paragraph(f"{activite['commune']}")

        doc.add_heading("Paroisse", level=2)
        doc.add_paragraph(f"{activite['paroisse']}")

        doc.add_heading(f"Unité physique", level=2)
        doc.add_paragraph(f"{activite['unite_physique']}")

        doc.add_heading(f"Quantité prévue", level=2)
        doc.add_paragraph(f"{activite['quantite_prevue']}")

        doc.add_heading(f"Cout de realisation", level=2)
        doc.add_paragraph(f"{activite['cout_realisation']}")

        doc.add_heading(f"Contribution des beneficiaires", level=2)
        doc.add_paragraph(f"{activite['contribution_beneficiaire']}")
        doc.add_heading(f"Contribution des partenaires", level=2)
        doc.add_paragraph(f"{activite['contribution_partenaire']}")
        doc.add_heading(f"Nombre des beneficiaires direct hommes", level=2)
        doc.add_paragraph(f"{activite['nbre_benef_direct_homme']}")
        doc.add_heading(f"Nombre des beneficiaires direct femmes", level=2)
        doc.add_paragraph(f"{activite['nbre_benef_direct_femme']}")
        doc.add_heading(f"Nombre total des beneficiaires direct", level=2)
        doc.add_paragraph(f"{activite['total_benef_direct']}")
        
        doc.add_heading(f"Partenaire financier principal", level=2)
        doc.add_paragraph(f"{activite['partenaires']}")

        doc.add_paragraph("_________________________________________________________________________________________________________")
    
    doc.add_paragraph()
    doc.add_heading("SYNTHESE REALISATION", level=1)

    for activite in activites_groupees_realisation:
        doc.add_heading(f"Projet : {activite['projet']}", level=2)

        doc.add_heading(f"Activitée", level=2)
        doc.add_paragraph(f"{activite['titre']}")

        doc.add_heading(f"Secteur", level=2)
        doc.add_paragraph(f"{activite['secteur']}")

        doc.add_heading(f"Domaine", level=2)
        doc.add_paragraph(f"{activite['domaine']}")

        doc.add_heading("Région", level=2)
        doc.add_paragraph(f"{activite['region']}")

        doc.add_heading("Province", level=2)
        doc.add_paragraph(f"{activite['province']}")

        doc.add_heading("Commune", level=2)
        doc.add_paragraph(f"{activite['commune']}")
        
        doc.add_heading("Paroisse", level=2)
        doc.add_paragraph(f"{activite['paroisse']}")

        doc.add_heading("Unité physique", level=2)
        doc.add_paragraph(f"{activite['unite_physique']}")

        doc.add_heading("Quantité réalisé", level=2)
        doc.add_paragraph(f"{activite['quantite_prevue']}")

        doc.add_heading("Période réalisé", level=2)
        doc.add_paragraph(f"Date de début du contrat (démarrage) : {activite['periode_prevue_debut']}")
        doc.add_paragraph(f"Date de fin du contrat : {activite['periode_prevue_fin']}")

        doc.add_heading("Responsable d'exécution", level=2)
        doc.add_paragraph(f"{activite['responsable']}")

        doc.add_heading("Coût de realisation", level=2)
        doc.add_paragraph(f"{activite['cout_realisation']}")

        doc.add_heading("Contribution bénéficiaire", level=2)
        doc.add_paragraph(f"{activite['contribution_beneficiaire']} FCFA")

        doc.add_heading("Contribution partenaire", level=2)
        doc.add_paragraph(f"{activite['contribution_partenaire']} FCFA")

        doc.add_heading("Nombre bénéficiare direct homme", level=2)
        doc.add_paragraph(f"{activite['nbre_benef_direct_homme']}")

        doc.add_heading("Nombre bénéficiare direct femme", level=2)
        doc.add_paragraph(f"{activite['nbre_benef_direct_femme']}")

        doc.add_heading("Nombre total de bénéficiaire direct", level=2)
        doc.add_paragraph(f"{activite['total_benef_direct']}")

        doc.add_heading(f"Partenaire financier principal", level=2)
        doc.add_paragraph(f"{activite['partenaireP']}")

        doc.add_paragraph("_________________________________________________________________________________________________________")
    
    doc.add_paragraph()
    doc.add_heading("SYNTHESE SITUATION FISCALE ET SOCIALE")

    for situation in situations:
        doc.add_heading(f"Utilisateur : { situation.utilisateur } ", level=2)
        doc.add_heading(f"Projet : { situation.id_projet.nom } ", level=2)

        doc.add_heading(f"SITUATIONS FISCALES ET SOCIALES DU SED AU 31 DECEMBRE", level=3)
        doc.add_paragraph(f"Impôts et taxes versées: { situation.impot } FCFA")
        doc.add_paragraph(f"Cotisations Sociales versées : { situation.cotisation } FCFA")
        doc.add_paragraph(f"Autres contributions fiscales versées : { situation.autre_contribution } FCFA")
        doc.add_paragraph(f"Total : { situation.total }")

        doc.add_heading(f"DERNIERS AUDITS COMPTABLES REALISES AU COURS DE L'ANNEE", level=3)
        audits = audits[situation.id]
        if audits:
            for p in audits:
                doc.add_paragraph(f"Désignation projets et programmes audités : {p['designation']} ")
                doc.add_paragraph(f"Date de réalisation : {p['date_realisation']} ")
                doc.add_paragraph(f"Nom du cabinet ayant conduit l'audit comptable : {p['nom_cabinet']} ")
                doc.add_paragraph("----------------------------------------")
        else:
            doc.add_paragraph(f"Désignation projets et programmes audités : {p['designation']} ")
            doc.add_paragraph(f"Date de réalisation : {p['date_realisation']} ")
            doc.add_paragraph(f"Nom du cabinet ayant conduit l'audit comptable : {p['nom_cabinet']} ")
        
        doc.add_paragraph("_________________________________________________________________________________________________________")

    doc.add_paragraph()
    doc.add_heading('DEPENSE DE FONCTIONNEMENT')

    for depense in depenses:
        doc.add_heading(f"Utilisateur : { depense.utilisateur }", level=2)
        doc.add_heading(f"Projet : { str(depense.id_projet.nom) }", level=2)

        doc.add_paragraph(f"Salaire et avantage du personnel: { depense.salaire } F CFA")
        doc.add_paragraph(f"Carburant et entretien des engins (motos, véhicules): { depense.carburant_entretien } FCFA")
        doc.add_paragraph(f"Frais de communication: { depense.frais_communication } F CFA")
        doc.add_paragraph(f"Acquisition de matériel (moto, ordinateurs,vélicule…) pour le fonctionnement : { depense.acquisition_materiel } F CFA")
        doc.add_paragraph(f"Rencontres de l'équipe projet (rencontres mensuelles et autres rencontres) : { depense.rencontre_equipe } F CFA")
        doc.add_paragraph(f"Contribution électricité et eau : { depense.contribution_electrique_eau } F CFA")
        doc.add_paragraph(f"Achat de consommables bureautiques : { depense.achat_consommable } F CFA")
        doc.add_paragraph(f"Entretien des locaux : { depense.entretien_locaux } F CFA")
        doc.add_paragraph(f"Maintenance des appareil : { depense.maintenance_appareil } F CFA")
        doc.add_paragraph(f"Expédition courrier : { depense.expedition_courier } F CFA")
        doc.add_paragraph(f"Frais bancaire : { depense.frais_bancaire } F CFA")
        doc.add_paragraph(f"Couts indirects : { depense.cout_indirect } F CFA")

        doc.add_paragraph("_________________________________________________________________________________________________________")


    output = BytesIO()
    doc.save(output)

    output.seek(0)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=globale_synthese.docx'
    response.write(output.getvalue())

    return response

# generer un excel pour la synthese globale
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_excel_globale(request):
    # Informations generales
    infos_generale = InfosGenerale.objects.all()
    partenariats_generale = {}
    objectifs_generale = {}

    for info in infos_generale:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_generale[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_generale[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]

    # Informations specifique
    infos_specifique = InfosSpecific.objects.all()
    objectif_specific = {}
    resultat_specific = {}

    for specific in infos_specifique:
        objectif_group = Objectif.objects.filter(id_infos_specifique=specific)
        objectif_specific[specific.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group ]

        resultat_group = Resultat.objects.filter(id_specific=specific)
        resultat_specific[specific.id] = [{
            'resultats': resultat.resultats
        } for resultat in resultat_group]

    # Planification
    activites_groupees_planification = []
    distinct_titres_planification = Activite.objects.values_list('titre', flat=True).distinct()

    for titre in distinct_titres_planification:
        activite_groupe = Activite.objects.filter(titre=titre)
        
        region = safe_join(activite_groupe.values_list('region', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        partenaires = safe_join(activite_groupe.values_list('partenaires', flat=True))

        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)
        
        activites_groupees_planification.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre': titre,
            'region': region,
            'province': province,
            'commune': commune,
            'paroisse': paroisse,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'total_benef_direct': total_benef_direct,
            'partenaires': partenaires
        })

    # Realisation
    distinct_titres_realisation = Realisation.objects.values_list('titre', flat=True).distinct()
    activites_groupees_realisation = []

    for titre in distinct_titres_realisation:
        activite_groupe = Realisation.objects.filter(titre=titre)

        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        region = safe_join(activite_groupe.values_list('region', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = safe_join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = safe_join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = safe_join(activite_groupe.values_list('responsable', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        partenaireP = safe_join(activite_groupe.values_list('partenaires', flat=True))

        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)

        activites_groupees_realisation.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre': titre,
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'partenaireP': partenaireP
        })

    # Situation
    situations = Situation.objects.all()

    audits = {}

    for situation in situations:
        audit_group = Audit.objects.filter(situation=situation)
        audits[situation.id] = [{
            'designation': audits.designation,
            'date_realisation': audits.date_realisation,
            'nom_cabinet': audits.nom_cabinet
        } for audits in audit_group]

    # Depense
    depenses = Depense.objects.all()

    wb = Workbook()
    ws = wb.active

    # Ajout des informations au fichier excel
    # Informations generales
    ws.append(['INFORMATIONS GENERALES'])
    headers = [
        'Créateur du projet', 'Nom de l\'organisation', 'Nature de l\'organisation', 'Sigle', 'Pays d\'origine', 'Région', 'Province', 'Commune', 'Ville/Secteur',
        'Boite postale', 'Numéro de téléphone fixe', 'Numéro de téléphone mobile', 'Adresse mail professionnelle', 'Site web', 'Nom et Prénom(s)', 'Nationalité',
        'Fonction(Président,...)', 'Numéro fixe', 'Numéro mobile', 'Dernier renouvèlement des Instances dirigeantes', 'Dernière Assemblée Générale Ordinaire',
        'Dernière session statutaire du bureau exécutif', 'Durée du mandat du bureau exécutif', 'Objectifs', 'Groupes cibles specifique', 'Hommes Nombre total du personnel', 'Femmes Nombre total du personnel', 'Hommes Employés nationaux Contrat à Durée Indéterminée (CDI)',
        'Femmes Employés nationaux Contrat à Durée Indéterminée (CDI)', 'Hommes Employés nationaux Contrat à Durée déterminée (CDD)', 'Femmes Employés nationaux Contrat à Durée déterminée (CDD)',
        'Hommes Employés expatriés Contrat à Durée Indéterminée (CDI)', 'Femmes Employés expatriés Contrat à Durée Indéterminée (CDI)', 'Hommes Employés expatriés Contrat à Durée déterminée (CDD)',
        'Femmes Employés expatriés Contrat à Durée déterminée (CDD)', 'Hommes Bénévoles ou volontaires Nationaux', 'Femmes Bénévoles ou volontaires Nationaux',
        'Hommes Bénévoles ou volontaires Expatriés', 'Femmes Bénévoles ou volontaires Expatriés', 'Hommes Personnel de l\'Administration publique en détachement', 'Femmes Personnel de l\'Administration publique en détachement',
        'Nom partenariats', 'N° de convention de partenariat / protocole d\'entente', 'Date de début d\'effet', 'Date de fin d\'effet'
    ]
    ws.append(headers)

    for info in infos_generale:
        partenariats = partenariats_generale[info.id]
        objectifs = objectifs_generale[info.id]

        objectifs_str = safe_join([o['objectifs'] for o in objectifs])

        if partenariats:
            for p in partenariats:
                row = [
                    info.utilisateur.username, info.nom_org , info.nature_org, info.sigle, info.pays_origine, info.region, info.province, info.commune,
                    info.village, info.boite_postale, info.numb_fixe, info.numb_mobile, info.adresse_mail, info.site_web, info.nom_complet_resp,
                    info.nationalite_resp, info.fonction_resp, info.numb_fixe_resp, info.numb_mobile_resp, info.renou_instance, info.assem_general,
                    info.session_statut, info.mandat_bureau,  objectifs_str, info.groupes_cibles, info.total_pers_homme, info.total_pers_femme, info.em_nation_cdi_homme, info.em_nation_cdi_femme,
                    info.em_nation_cdd_homme, info.em_nation_cdd_femme, info.em_expa_cdi_homme, info.em_expa_cdi_femme, info.em_expa_cdd_homme,
                    info.em_expa_cdd_femme, info.benevol_nation_homme, info.benevol_nation_femme, info.benevol_expa_homme, info.benevol_expa_femme,
                    info.personnel_admin_homme, info.personnel_admin_femme, p['nom'], p['numero'], p['date_debut'], p['date_fin']
                ]
                ws.append(row)
        else:
            row = [
                info.utilisateur.username, info.nom_org, info.nature_org, info.sigle, info.pays_origine, info.region, info.province, info.commune,
                    info.village, info.boite_postale, info.numb_fixe, info.numb_mobile, info.adresse_mail, info.site_web, info.nom_complet_resp,
                    info.nationalite_resp, info.fonction_resp, info.numb_fixe_resp, info.numb_mobile_resp, info.renou_instance, info.assem_general,
                    info.session_statut, info.mandat_bureau,  objectifs_str, info.groupes_cibles, info.total_pers_homme, info.total_pers_femme, info.em_nation_cdi_homme, info.em_nation_cdi_femme,
                    info.em_nation_cdd_homme, info.em_nation_cdd_femme, info.em_expa_cdi_homme, info.em_expa_cdi_femme, info.em_expa_cdd_homme,
                    info.em_expa_cdd_femme, info.benevol_nation_homme, info.benevol_nation_femme, info.benevol_expa_homme, info.benevol_expa_femme,
                    info.personnel_admin_homme, info.personnel_admin_femme, '', '', '', ''
            ]
            ws.append(row)

    ws.append([])

    # Informations specifique
    ws.append(['INFORMATIONS SPECIFIQUES'])
    headers_specific = [
        'Nom du projet', 'Date de debut du projet', 'Date de fin projet', 'Cout', 'Nombre de beneficiaire directs hommes',
        'Nombre de beneficiaire direct femmes', 'Objectif principale', 'Objectifs secondaires', 'Resultats', 'Partenaire financier principal'
    ]
    ws.append(headers_specific)

    for specific in infos_specifique:
        objectifs = objectif_specific[specific.id]
        resultats = resultat_specific[specific.id]

        objectifs_str = ', '.join([o['objectifs'] for o in objectifs]) + '\n'
        resultats_str = ', '.join([o['resultats'] for o in resultats]) + '\n'

        row_specific = [
            specific.nom, specific.date_debut, specific.date_fin, specific.cout, specific.benef_direct_homme,
            specific.benef_direct_femme, specific.objectifs_principals, objectifs_str, resultats_str, specific.partenaires
        ]
        ws.append(row_specific)

    ws.append([])

    # Planification
    ws.append(['SYNTHESE PLANIFICATION'])
    headers_planification = ['Projet', 'Titre', 'Secteur', 'Domaine', 'Region', 'Province', 'Commune', 'Paroisse', 'Unité Physique', 
               'Quantité Prévue', 'Cout de realisation', 'Contribution des beneficiaires', 
               'Contribution des partenaires', 'Nombre des beneficiaires direct hommes', 'Nombre des beneficiaires direct femmes', 'Nombre total des beneficiaires direct', 
               'Partenaire financier principal']
    ws.append(headers_planification)

    for activite in activites_groupees_planification:
        projet = str(activite['projet'])
        titre = str(activite['titre'])
        secteur = str(activite['secteur'])
        domaine = str(activite['domaine'])
        region = str(activite['region'])
        province = str(activite['province'])
        commune = str(activite['commune'])
        paroisse = str(activite['paroisse'])
        unite_physique = activite['unite_physique']
        quantite_prevue = activite['quantite_prevue']
        cout_realisation = activite['cout_realisation']
        contribution_beneficiaire = activite['contribution_beneficiaire']
        contribution_partenaire = activite['contribution_partenaire']
        nbre_benef_direct_homme = activite['nbre_benef_direct_homme']
        nbre_benef_direct_femme = activite['nbre_benef_direct_femme']
        total_benef_direct = activite['total_benef_direct']
        partenaires = str(activite['partenaires'])
        
        row_planification = [projet, titre, secteur, domaine, region, province, commune, paroisse, unite_physique, 
                   quantite_prevue, cout_realisation, contribution_beneficiaire, 
                   contribution_partenaire, nbre_benef_direct_homme, nbre_benef_direct_femme, total_benef_direct, 
                   partenaires]
        ws.append(row_planification)
    
    ws.append([])

    # Realisation
    ws.append(['SYNTHESE REALISATION'])
    headers_realisation = ['Projet', 'Titre', 'Secteur', 'Domaine', 'Région', 'Province', 'Commune', 'Paroisse', 'Unité physique', 'Quantité réalisé', 'Date de début du contrat(démarrage)', 'Date de fin du contrat', 'Responsable d\'exécution', 'Coût de réalisation', 'Contribution bénéficiaire', 'Contribution partenaire', 'Nombre bénéficiare direct homme', 'Nombre bénéficiare direct femme', 'Nombre total de bénéficiaire direct', 
                           'Partenaire financier principal']
    ws.append(headers_realisation)

    for activite in activites_groupees_realisation:
        projet = str(activite['projet'])
        titre = str(activite['titre'])
        secteur = str(activite['secteur'])
        domaine = str(activite['domaine'])
        commune = str(activite['commune'])
        province = str(activite['province'])
        region = str(activite['region'])
        paroisse = str(activite['paroisse'])
        unite_physique = activite['unite_physique']
        quantite_realise = activite['quantite_prevue']
        periode_prevue_debut = activite['periode_prevue_debut']
        periode_prevue_fin = activite['periode_prevue_fin']
        responsable = str(activite['responsable'])
        cout_realisation = activite['cout_realisation']
        contribution_beneficiaire = activite['contribution_beneficiaire']
        contribution_partenaire  = activite['contribution_partenaire']
        total_benef_direct = activite['total_benef_direct']
        nbre_benef_direct_homme = activite['nbre_benef_direct_homme']
        nbre_benef_direct_femme = activite['nbre_benef_direct_femme']
        partenaireP = str(activite['partenaireP'])

        row_realisation = [projet, titre, secteur, domaine, region, province, commune, paroisse, unite_physique, quantite_realise, periode_prevue_debut, periode_prevue_fin, responsable, cout_realisation, contribution_beneficiaire, contribution_partenaire, nbre_benef_direct_homme, nbre_benef_direct_femme, total_benef_direct, 
                           partenaireP]
        ws.append(row_realisation)
    
    ws.append([])

    # Situation fiscale et sociale
    ws.append(['SYNTHESE SITUATION FISCALE ET SOCIALE'])
    headers_situation = [
        'Utilisateur', 'Projet', 'Impôts et taxes versées', 'Cotisations Sociales versées', 'Autres contributions fiscales versées', 'Total',
        'Désignation projets et programmes audités', 'Date de réalisation', 'Nom du cabinet ayant conduit l\'audit comptable'
    ]
    ws.append(headers_situation)

    for situation in situations:
        audits = audits.get(situation.id, [])

        if audits:
            designation = safe_join([audit['designation'] for audit in audits])
            date_realisation = safe_join([audit['date_realisation'] for audit in audits])
            nom_cabinet = safe_join([audit['nom_cabinet'] for audit in audits])
        else:
            designation = ''
            date_realisation = ''
            nom_cabinet = ''

        row_situations = [
            situation.utilisateur, situation.id_projet.nom, situation.impot, situation.cotisation, situation.autre_contribution, situation.total,
            designation, date_realisation, nom_cabinet
        ]
        ws.append(row_situations)
    
    ws.append([])

    # Depense de fonctionnement
    ws.append(['SYNTHESE DE FONCTIONNEMENT'])
    headers_depense = [
        'Utilisateur', 'Projet', 'Salaire et avantage du personnel (F CFA)', 'Carburant et entretien des engins (motos, véhicules) (F CFA)',
        'Frais de communication (F CFA)', 'Acquisition de matériel (moto, ordinateurs,vélicule…) pour le fonctionnement (F CFA)',
        'Rencontres de l\'équipe projet (rencontres mensuelles et autres rencontres) (F CFA)', 'Contribution électricité et eau (F CFA)',
        'Achat de consommables bureautiques (F CFA)', 'Entretien des locaux (F CFA)', 'Maintenance des appareil (F CFA)',
        'Expédition courrier (F CFA)', 'Frais bancaire (F CFA)', 'Couts indirects (F CFA)'
    ]
    ws.append(headers_depense)

    for depense in depenses:
        row_depenses = [
            str(depense.utilisateur), str(depense.id_projet.nom), depense.salaire, depense.carburant_entretien,
            depense.frais_communication, depense.acquisition_materiel, depense.rencontre_equipe, depense.contribution_electrique_eau,
            depense.achat_consommable, depense.entretien_locaux, depense.maintenance_appareil, depense.expedition_courier,
            depense.frais_bancaire, depense.cout_indirect
        ]
        ws.append(row_depenses)


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=globale_synthese.xlsx'

    wb.save(response)

    return response

# generer un PDF pour les depenses
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_pdf_depense(request):
    depenses = Depense.objects.all()

    depense_dict = {}
    for depense in depenses:
        autre = AutreDepense.objects.filter(depense = depense)
        depense_dict[depense.id] = [{
            'intitule': depens.intitule,
            'prix': depens.prix
        } for depens in autre]

    context = {
        'depenses': depenses,
        'depense_dict': depense_dict
    }
    template_path = 'service/invoice_depense.html'
    template = get_template(template_path)
    html = template.render(context)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="synthese_depense.pdf"'
    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la generation du PDF')
    return response

# generer un word pour les depenses
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_word_depense(request):
    depenses = Depense.objects.all()

    doc = Document()
    doc.add_heading('SYNTHÈSE GLOBALE DEPENSE DE FONCTIONNEMENT', level=1)
    
    for depense in depenses:
        doc.add_heading(f"Utilisateur : { depense.utilisateur }", level=2)
        doc.add_heading(f"Projet : { str(depense.id_projet.nom) }", level=2)

        doc.add_paragraph(f"Salaire et avantage du personnel: { depense.salaire } F CFA")
        doc.add_paragraph(f"Carburant et entretien des engins (motos, véhicules): { depense.carburant_entretien } FCFA")
        doc.add_paragraph(f"Frais de communication: { depense.frais_communication } F CFA")
        doc.add_paragraph(f"Acquisition de matériel (moto, ordinateurs,vélicule…) pour le fonctionnement : { depense.acquisition_materiel } F CFA")
        doc.add_paragraph(f"Rencontres de l'équipe projet (rencontres mensuelles et autres rencontres) : { depense.rencontre_equipe } F CFA")
        doc.add_paragraph(f"Contribution électricité et eau : { depense.contribution_electrique_eau } F CFA")
        doc.add_paragraph(f"Achat de consommables bureautiques : { depense.achat_consommable } F CFA")
        doc.add_paragraph(f"Entretien des locaux : { depense.entretien_locaux } F CFA")
        doc.add_paragraph(f"Maintenance des appareil : { depense.maintenance_appareil } F CFA")
        doc.add_paragraph(f"Expédition courrier : { depense.expedition_courier } F CFA")
        doc.add_paragraph(f"Frais bancaire : { depense.frais_bancaire } F CFA")
        doc.add_paragraph(f"Couts indirects : { depense.cout_indirect } F CFA")

        doc.add_paragraph("_________________________________________________________________________________________________________")
    
    output = BytesIO()
    doc.save(output)

    output.seek(0)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=synthese_depense.docx'
    response.write(output.getvalue())

    return response

# generer un excel des depenses
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_excel_depense(request):
    depenses = Depense.objects.all()

    wb = Workbook()
    ws = wb.active

    headers = [
        'Utilisateur', 'Projet', 'Salaire et avantage du personnel (F CFA)', 'Carburant et entretien des engins (motos, véhicules) (F CFA)',
        'Frais de communication (F CFA)', 'Acquisition de matériel (moto, ordinateurs,vélicule…) pour le fonctionnement (F CFA)',
        'Rencontres de l\'équipe projet (rencontres mensuelles et autres rencontres) (F CFA)', 'Contribution électricité et eau (F CFA)',
        'Achat de consommables bureautiques (F CFA)', 'Entretien des locaux (F CFA)', 'Maintenance des appareil (F CFA)',
        'Expédition courrier (F CFA)', 'Frais bancaire (F CFA)', 'Couts indirects (F CFA)'
    ]
    ws.append(headers)

    for depense in depenses:
        row = [
            str(depense.utilisateur), str(depense.id_projet.nom), depense.salaire, depense.carburant_entretien,
            depense.frais_communication, depense.acquisition_materiel, depense.rencontre_equipe, depense.contribution_electrique_eau,
            depense.achat_consommable, depense.entretien_locaux, depense.maintenance_appareil, depense.expedition_courier,
            depense.frais_bancaire, depense.cout_indirect
        ]
        ws.append(row)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=synthese_depense.xlsx'

    wb.save(response)

    return response

# Generer le pdf de synthese planfication
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_pdf_planification(request):
    distinct_titres = Activite.objects.values_list('titre', flat=True).distinct()
    nom_final = []
    activites_groupees = []
    for titre in distinct_titres:
        activite_groupe = Activite.objects.filter(titre=titre)
        for activite in activite_groupe:
            projet = activite.id_projet
            nom_final.append(projet.nom)
        
        region = safe_join(activite_groupe.values_list('region', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        partenaires = safe_join(activite_groupe.values_list('partenaires', flat=True))
        
        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)

        activites_groupees.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre':titre,
            'region': region,
            'province': province,
            'commune': commune,
            'paroisse': paroisse,
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'total_benef_direct': total_benef_direct,
            'partenaires': partenaires
        })

    # Charger le template HTML
    template_path = 'service/invoice_planification.html'
    template = get_template(template_path)
    context = {
        'activites_groupees': activites_groupees
    }
    html = template.render(context)
    
    # Créer un objet HttpResponse avec le type de contenu PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="synthese_planification.pdf"'
    
    # Convertir le template HTML en PDF
    pisa_status = pisa.CreatePDF(html, dest=response)
    
    # Si la conversion a réussi, retourner la réponse avec le PDF généré
    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la génération du PDF')
    return response

# Generer le word de planification
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_word_planification(request):
    activites_groupees = []
    nom_final = []
    distinct_titres = Activite.objects.values_list('titre', flat=True).distinct()

    for titre in distinct_titres:
        activite_groupe = Activite.objects.filter(titre=titre)
        for activite in activite_groupe:
            projet = activite.id_projet
            nom_final.append(projet.nom)

        region = safe_join(activite_groupe.values_list('region', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        partenaires = safe_join(activite_groupe.values_list('partenaires', flat=True))

        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)
        
        activites_groupees.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre': titre,
            'region': region,
            'province': province,
            'commune': commune,
            'paroisse': paroisse,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'total_benef_direct': total_benef_direct,
            'partenaires': partenaires          
        })

    # creer un document word
    doc = Document()

    # ajouter le titre
    doc.add_heading('SYNTHÈSE GLOBALE PLANIFICATION OPÉRATIONNELLE', level=1)

    for activite in activites_groupees:
        doc.add_heading(f"{activite['projet']}")
        
        doc.add_heading(f"Activitée", level=2)
        doc.add_paragraph(f"{activite['titre']}")

        doc.add_heading(f"Secteur", level=2)
        doc.add_paragraph(f"{activite['secteur']}")

        doc.add_heading(f"Domaine", level=2)
        doc.add_paragraph(f"{activite['domaine']}")

        doc.add_heading(f"Région", level=2)
        doc.add_paragraph(f"{activite['region']}")

        doc.add_heading(f"Province", level=2)
        doc.add_paragraph(f"{activite['province']}")

        doc.add_heading(f"Commune", level=2)
        doc.add_paragraph(f"{activite['commune']}")

        doc.add_heading("Paroisse", level=2)
        doc.add_paragraph(f"{activite['paroisse']}")

        doc.add_heading(f"Unité physique", level=2)
        doc.add_paragraph(f"{activite['unite_physique']}")

        doc.add_heading(f"Quantité prévue", level=2)
        doc.add_paragraph(f"{activite['quantite_prevue']}")

        doc.add_heading(f"Cout de realisation", level=2)
        doc.add_paragraph(f"{activite['cout_realisation']}")

        doc.add_heading(f"Contribution des beneficiaires", level=2)
        doc.add_paragraph(f"{activite['contribution_beneficiaire']}")
        doc.add_heading(f"Contribution des partenaires", level=2)
        doc.add_paragraph(f"{activite['contribution_partenaire']}")
        doc.add_heading(f"Nombre des beneficiaires direct hommes", level=2)
        doc.add_paragraph(f"{activite['nbre_benef_direct_homme']}")
        doc.add_heading(f"Nombre des beneficiaires direct femmes", level=2)
        doc.add_paragraph(f"{activite['nbre_benef_direct_femme']}")
        doc.add_heading(f"Nombre total des beneficiaires direct", level=2)
        doc.add_paragraph(f"{activite['total_benef_direct']}")
        
        doc.add_heading(f"Partenaire financier principal", level=2)
        doc.add_paragraph(f"{activite['partenaires']}")

        doc.add_paragraph("_________________________________________________________________________________________________________")

    # Créer un flux de mémoire pour stocker le document Word
    output = BytesIO()
    doc.save(output)

    # Réinitialiser le flux à la position de départ
    output.seek(0)

    # Creer une reponse Http avec le contenu du document word
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=synthese_planification.docx'
    response.write(output.getvalue())

    return response    

# Generer le excel de planification
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_excel_planification(request):
    activites_groupees = []
    nom_final = []
    distinct_titres = Activite.objects.values_list('titre', flat=True).distinct()

    for titre in distinct_titres:
        activite_groupe = Activite.objects.filter(titre=titre)
        for activite in activite_groupe:
            projet = activite.id_projet
            nom_final.append(projet.nom)
        
        region = safe_join(activite_groupe.values_list('region', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        partenaires = safe_join(activite_groupe.values_list('partenaires', flat=True))

        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)
        
        activites_groupees.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre': titre,
            'region': region,
            'province': province,
            'commune': commune,
            'paroisse': paroisse,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'total_benef_direct': total_benef_direct,
            'partenaires': partenaires
        })

    # Créer un nouveau classeur Excel
    wb = Workbook()
    # Créer une nouvelle feuille dans le classeur
    ws = wb.active
    # Ajouter des en-têtes de colonnes
    ws.append(['Projet', 'Titre', 'Secteur', 'Domaine', 'Region', 'Province', 'Commune', 'Paroisse', 'Unité Physique', 
               'Quantité Prévue', 'Cout de realisation', 'Contribution des beneficiaires', 
               'Contribution des partenaires', 'Nombre des beneficiaires direct hommes', 'Nombre des beneficiaires direct femmes', 'Nombre total des beneficiaires direct', 'Partenaire financier principal'])
    # Ajouter les données à la feuille Excel
    for activite in activites_groupees:
        projet = str(activite['projet'])
        titre = str(activite['titre'])
        secteur = str(activite['secteur'])
        domaine = str(activite['domaine'])
        region = str(activite['region'])
        province = str(activite['province'])
        commune = str(activite['commune'])
        paroisse = str(activite['paroisse'])
        unite_physique = activite['unite_physique']
        quantite_prevue = activite['quantite_prevue']
        cout_realisation = activite['cout_realisation']
        contribution_beneficiaire = activite['contribution_beneficiaire']
        contribution_partenaire = activite['contribution_partenaire']
        nbre_benef_direct_homme = activite['nbre_benef_direct_homme']
        nbre_benef_direct_femme = activite['nbre_benef_direct_femme']
        total_benef_direct = activite['total_benef_direct']
        partenaires = str(activite['partenaires'])
        

        ws.append([projet, titre, secteur, domaine, region, province, commune, paroisse, unite_physique, 
                   quantite_prevue, cout_realisation, contribution_beneficiaire, 
                   contribution_partenaire, nbre_benef_direct_homme, nbre_benef_direct_femme, total_benef_direct, partenaires ])
    
    # Créer une réponse HTTP pour le fichier Excel
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=synthese_planification.xlsx'

    # Enregistrer le classeur Excel dans la réponse HTTP
    wb.save(response)

    return response

# Generer pdf suivi
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_pdf_suivi(request):
    distinct_titres = Realisation.objects.values_list('titre', flat=True).distinct()
    nom_final = []
    secteur = []
    domaine = []
    ##################
    activites_groupees = []
    for titre in distinct_titres:
        #test_activite = Realisation.objects.get(titre=titre)
        activite_groupe = Realisation.objects.filter(titre=titre)
        for activite in activite_groupe:
            projet = activite.id_projet
            secteur_test = activite.id_secteur
            domaine_test = activite.id_sous_secteur
            nom_final.append(projet.nom)
            secteur.append(secteur_test)
            domaine.append(domaine_test)
    
        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        region = safe_join(activite_groupe.values_list('region', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = safe_join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = safe_join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = safe_join(activite_groupe.values_list('responsable', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        partenaireP = safe_join(activite_groupe.values_list('partenaires', flat=True))

        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)

        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_realisation=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre': titre,
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'partenaireP': partenaireP,
            'partenaires': list(partenaires_groupes.values())
        })
    template_path = 'service/invoice_realisation.html'
    template = get_template(template_path)
    context = {
        'activites_groupees': activites_groupees
    }
    html = template.render(context)

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="synthese_suivi.pdf"'

    pisa_status = pisa.CreatePDF(html, dest=response)

    if pisa_status.err:
        return HttpResponse('Une erreur est survenue lors de la génération du PDF')
    return response

# Generer le word de suivi
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_word_suivi(request):
    distinct_titres = Realisation.objects.values_list('titre', flat=True).distinct()
    ##################
    activites_groupees = []
    nom_final = []
    secteur = []
    domaine = []
    for titre in distinct_titres:
        #test_activite = Realisation.objects.get(titre=titre)
        activite_groupe = Realisation.objects.filter(titre=titre)
        for activite in activite_groupe:
            projet = activite.id_projet
            secteur_test = activite.id_secteur
            domaine_test = activite.id_sous_secteur
            nom_final.append(projet.nom)
            secteur.append(secteur_test)
            domaine.append(domaine_test)

        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        region = safe_join(activite_groupe.values_list('region', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = safe_join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = safe_join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = safe_join(activite_groupe.values_list('responsable', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        partenaireP = safe_join(activite_groupe.values_list('partenaires', flat=True))
        
        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)
        
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_realisation=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre': titre,
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'partenaireP': partenaireP,
            'partenaires': list(partenaires_groupes.values())
            })
    
    doc = Document()

    doc.add_heading("SYNTHÈSE GLOBALE SUIVI DES ACTIVITÉES", level=1)

    for activite in activites_groupees:
        doc.add_heading(f"Projet : {activite['projet']}", level=2)

        doc.add_heading(f"Activitée", level=2)
        doc.add_paragraph(f"{activite['titre']}")

        doc.add_heading(f"Secteur", level=2)
        doc.add_paragraph(f"{activite['secteur']}")

        doc.add_heading(f"Domaine", level=2)
        doc.add_paragraph(f"{activite['domaine']}")

        doc.add_heading("Région", level=2)
        doc.add_paragraph(f"{activite['region']}")

        doc.add_heading("Province", level=2)
        doc.add_paragraph(f"{activite['province']}")

        doc.add_heading("Commune", level=2)
        doc.add_paragraph(f"{activite['commune']}")
        
        doc.add_heading("Paroisse", level=2)
        doc.add_paragraph(f"{activite['paroisse']}")

        doc.add_heading("Unité physique", level=2)
        doc.add_paragraph(f"{activite['unite_physique']}")

        doc.add_heading("Quantité réalisé", level=2)
        doc.add_paragraph(f"{activite['quantite_prevue']}")

        doc.add_heading("Période réalisé", level=2)
        doc.add_paragraph(f"Date de début du contrat (démarrage) : {activite['periode_prevue_debut']}")
        doc.add_paragraph(f"Date de fin du contrat : {activite['periode_prevue_fin']}")

        doc.add_heading("Responsable d'exécution", level=2)
        doc.add_paragraph(f"{activite['responsable']}")

        doc.add_heading("Coût de realisation", level=2)
        doc.add_paragraph(f"{activite['cout_realisation']}")

        doc.add_heading("Contribution bénéficiaire", level=2)
        doc.add_paragraph(f"{activite['contribution_beneficiaire']} FCFA")

        doc.add_heading("Contribution partenaire", level=2)
        doc.add_paragraph(f"{activite['contribution_partenaire']} FCFA")

        doc.add_heading("Nombre bénéficiare direct homme", level=2)
        doc.add_paragraph(f"{activite['nbre_benef_direct_homme']}")

        doc.add_heading("Nombre bénéficiare direct femme", level=2)
        doc.add_paragraph(f"{activite['nbre_benef_direct_femme']}")

        doc.add_heading("Nombre total de bénéficiaire direct", level=2)
        doc.add_paragraph(f"{activite['total_benef_direct']}")

        doc.add_heading(f"Partenaire financier principal", level=2)
        doc.add_paragraph(f"{activite['partenaireP']}")

        # doc.add_heading("Partenaires financiers", level=2)
        # for partenaire in activite['partenaires']:
        #     doc.add_paragraph(f"Nom du partenaire : {partenaire['nom']}")
        #     doc.add_paragraph(f"Part du partenaire : {partenaire['part']}")
        #     doc.add_paragraph("----------------------------------------")
    
    doc.add_paragraph("_________________________________________________________________________________________________________")

    output = BytesIO()
    doc.save(output)

    output.seek(0)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=synthese_suivi.docx'
    response.write(output.getvalue())

    return response

# Generer le excel de suivi
@login_required(login_url='/login/')
@responsable_charger_projet_required
def generate_excel_suivi(request):
    distinct_titres = Realisation.objects.values_list('titre', flat=True).distinct()
    ##################
    activites_groupees = []
    nom_final = []

    for titre in distinct_titres:
        activite_groupe = Realisation.objects.filter(titre=titre)
        for activite in activite_groupe:
            projet = activite.id_projet
            nom_final.append(projet.nom)

        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        region = safe_join(activite_groupe.values_list('region', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = safe_join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = safe_join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = safe_join(activite_groupe.values_list('responsable', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        partenaireP = safe_join(activite_groupe.values_list('partenaires', flat=True))

        id_projets = activite_groupe.values_list('id_projet', flat=True)
        id_secteurs = activite_groupe.values_list('id_secteur', flat=True)
        id_domaines = activite_groupe.values_list('id_sous_secteur', flat=True)

        projets = Projet.objects.filter(id__in=id_projets)
        noms_projets = [projet.nom for projet in projets]
        projetTest = ' | '.join(noms_projets)

        secteurs = Secteur.objects.filter(id__in=id_secteurs)
        noms_secteurs = [secteur.titre for secteur in secteurs]
        secteurTest = ' | '.join(noms_secteurs)

        domaines = SousSecteur.objects.filter(id__in=id_domaines)
        noms_domaines = [domaine.titre for domaine in domaines]
        domaineTest = ' | '.join(noms_domaines)
        
        
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_realisation=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'projet': projetTest,
            'secteur': secteurTest,
            'domaine': domaineTest,
            'titre': titre,
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'partenaireP': partenaireP,
            'partenaires': list(partenaires_groupes.values())
        })
    
    wb = Workbook()
    ws = wb.active

    ws.append(['Projet', 'Titre', 'Secteur', 'Domaine', 'Région', 'Province', 'Commune', 'Paroisse', 'Unité physique', 'Quantité réalisé', 'Date de début du contrat(démarrage)', 'Date de fin du contrat', 'Responsable d\'exécution', 'Coût de réalisation', 'Contribution bénéficiaire', 'Contribution partenaire', 'Nombre bénéficiare direct homme', 'Nombre bénéficiare direct femme', 'Nombre total de bénéficiaire direct', 'Partenaire financier principal'])
    for activite in activites_groupees:
        projet = str(activite['projet'])
        titre = str(activite['titre'])
        secteur = str(activite['secteur'])
        domaine = str(activite['domaine'])
        commune = str(activite['commune'])
        province = str(activite['province'])
        region = str(activite['region'])
        paroisse = str(activite['paroisse'])
        unite_physique = activite['unite_physique']
        quantite_realise = activite['quantite_prevue']
        periode_prevue_debut = activite['periode_prevue_debut']
        periode_prevue_fin = activite['periode_prevue_fin']
        responsable = str(activite['responsable'])
        cout_realisation = activite['cout_realisation']
        contribution_beneficiaire = activite['contribution_beneficiaire']
        contribution_partenaire  = activite['contribution_partenaire']
        total_benef_direct = activite['total_benef_direct']
        nbre_benef_direct_homme = activite['nbre_benef_direct_homme']
        nbre_benef_direct_femme = activite['nbre_benef_direct_femme']
        partenaireP = str(activite['partenaireP'])
        nom_partenaire = ', '.join([f"{partenaire['nom']}" for partenaire in activite['partenaires']])
        part_partenaire = ', '.join([f"{partenaire['part']}" for partenaire in activite['partenaires']])

        ws.append([projet, titre, secteur, domaine, region, province, commune, paroisse, unite_physique, quantite_realise, periode_prevue_debut, periode_prevue_fin, responsable, cout_realisation, contribution_beneficiaire, contribution_partenaire, nbre_benef_direct_homme, nbre_benef_direct_femme, total_benef_direct, partenaireP])
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=synthese_suivi.xlsx'

    wb.save(response)

    return response

@login_required(login_url='/login/')
@responsable_charger_projet_required
def globale_generale(request):
    infos_test = InfosGenerale.objects.all()

    partenariats_dict = {}
    objectifs_dict = {}

    for info in infos_test:
        # Récupérer les partenariats associés
        partenariat_group = Partenariat.objects.filter(id_general=info)
        partenariats_dict[info.id] = [{
            'nom': partenariat.nom,
            'numero': partenariat.numero,
            'date_debut': partenariat.date_debut,
            'date_fin': partenariat.date_fin
        } for partenariat in partenariat_group]

        # Récupérer les objectifs associés
        objectif_group = Objectif.objects.filter(id_general=info)
        objectifs_dict[info.id] = [{
            'objectifs': objectif.objectifs
        } for objectif in objectif_group]

    context = {
        'infos_test': infos_test,
        'partenariats_dict': partenariats_dict,
        'objectifs_dict': objectifs_dict
    }

    return render(request, 'service/globale_generale.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def globale_situation(request):
    situations = Situation.objects.all()
    audit = {}

    for situation in situations:
        audit_group = Audit.objects.filter(situation=situation)
        audit[situation.id] = [{
            'designation': audits.designation,
            'date_realisation': audits.date_realisation,
            'nom_cabinet': audits.nom_cabinet
        } for audits in audit_group]

    context = {
        'situations': situations,
        'audit': audit
    }
    return render(request, 'service/globale_situation.html', context)

@login_required(login_url='/login/')
@responsable_required
def globale_planification(request):
    #
    distinct_titres = Activite.objects.values_list('titre', flat=True).distinct()
    ##################
    activites_groupees = []
    for titre in distinct_titres:
        activite_groupe = Activite.objects.filter(titre=titre)

        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        region = safe_join(activite_groupe.values_list('region', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))

        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))

        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        partenaireP = safe_join(activite_groupe.values_list('partenaires', flat=True))
             
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_activite=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'partenaireP': partenaireP,
            'partenaires': list(partenaires_groupes.values())
        })
    context = {
        'projets': 'projets',
        'activites_groupees': activites_groupees,
        'activieplus': 'activiteplus'
    }
    return render(request, 'service/globale.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def globale_suivi(request):
    #
    distinct_titres = Realisation.objects.values_list('titre', flat=True).distinct()
    ##################
    activites_groupees = []
    for titre in distinct_titres:
        activite_groupe = Realisation.objects.filter(titre=titre)
        
        commune = safe_join(activite_groupe.values_list('commune', flat=True))
        province = safe_join(activite_groupe.values_list('province', flat=True))
        region = safe_join(activite_groupe.values_list('region', flat=True))
        paroisse = safe_join(activite_groupe.values_list('paroisse', flat=True))
        unite_physique = safe_join(activite_groupe.values_list('unite_physique', flat=True))
        quantite_prevue = safe_sum(activite_groupe.values_list('quantite_prevue', flat=True))
        periode_prevue_debut = safe_join(activite_groupe.values_list('periode_prevue_debut', flat=True))
        periode_prevue_fin = safe_join(activite_groupe.values_list('periode_prevue_fin', flat=True))
        responsable = safe_join(activite_groupe.values_list('responsable', flat=True))
        cout_realisation = safe_sum(activite_groupe.values_list('cout_realisation', flat=True))
        contribution_beneficiaire = safe_sum(activite_groupe.values_list('contribution_beneficiaire', flat=True))
        contribution_partenaire = safe_sum(activite_groupe.values_list('contribution_partenaire', flat=True))
        total_benef_direct = safe_sum(activite_groupe.values_list('total_benef_direct', flat=True))
        nbre_benef_direct_homme = safe_sum(activite_groupe.values_list('nbre_benef_direct_homme', flat=True))
        nbre_benef_direct_femme = safe_sum(activite_groupe.values_list('nbre_benef_direct_femme', flat=True))
        partenaireP = safe_join(activite_groupe.values_list('partenaires', flat=True))
        
        
        partenaires = []
        for activite in activite_groupe:
            partenaires_activite = Partenaire.objects.filter(id_realisation=activite)
            for partenaire in partenaires_activite:
                partenaires.append({
                    'nom': partenaire.nom,
                    'part': partenaire.part
                })

        partenaires_groupes = {}
        for partenaire in partenaires:
            nom = partenaire['nom']
            part = partenaire['part']
            if nom in partenaires_groupes:
                partenaires_groupes[nom]['part'] += part
            else:
                partenaires_groupes[nom] = {'nom': nom, 'part': part}

        activites_groupees.append({
            'commune': commune,
            'province': province,
            'region': region,
            'paroisse': paroisse,
            'titre': titre,
            'unite_physique': unite_physique,
            'quantite_prevue': quantite_prevue,
            'periode_prevue_debut': periode_prevue_debut,
            'periode_prevue_fin': periode_prevue_fin,
            'responsable': responsable,
            'cout_realisation': cout_realisation,
            'contribution_beneficiaire': contribution_beneficiaire,
            'contribution_partenaire': contribution_partenaire,
            'total_benef_direct': total_benef_direct,
            'nbre_benef_direct_homme': nbre_benef_direct_homme,
            'nbre_benef_direct_femme': nbre_benef_direct_femme,
            'partenaireP': partenaireP,
            'partenaires': list(partenaires_groupes.values())
        })
    context = {
        'projets': 'projets',
        'activites_groupees': activites_groupees,
        'activieplus': 'activiteplus'
    }
    return render(request, 'service/globale1.html', context)


@login_required(login_url='/login/')
@all_user_required
def index(request):
    user = request.user
    total_service = Projet.objects.count()
    total_service_user = Projet.objects.filter(utilisateur=user).count()
    total_planifier = Activite.objects.count()
    total_realisation = Realisation.objects.count()
    total_user = User.objects.count()

    projets = Projet.objects.annotate(
        activites_count=Coalesce(Count('id_projet', distinct=True), 0),
        realisations_count=Coalesce(Count('id_projet_plus', distinct=True), 0),
        taux_realisation=Case(
            When(activites_count__gt=0, 
                 then=F('realisations_count') * 100.0 / F('activites_count')),
            default=0,
            output_field=IntegerField(),
        )
    ).filter(activites_count__gt=0)

    projets_data = []
    for projet in projets:
        projets_data.append({
            'nom': projet.nom,
            'taux': float(projet.taux_realisation)
        })

    # Pour le graphique des budgets
    infos_budgets = InfosSpecific.objects.filter(
        budget__gt=0
    ).values('nom', 'budget', 'depense_globale')

    budgets_data = []
    for info in infos_budgets:
        budgets_data.append({
            'nom': info['nom'],
            'budget_total': float(info['budget'] or 0),
            'depenses_total': float(info['depense_globale'] or 0)
        })


    context = {
        'projets_data': projets_data,
        'budgets_data': budgets_data,
        'total_service' : total_service,
        'user': user,
        'total_service_user': total_service_user,
        'total_planifier': total_planifier,
        'total_realisation': total_realisation,
        'total_user': total_user

    }
    
    return render(request, 'service/dashboard.html', context)

@login_required(login_url='/login/')
@all_user_required
def service(request):
    user = request.user
    service = Projet.objects.all()
    service_plus = Projet.objects.filter(utilisateur=user)
    context = {
        'service' : service
    }

    return render(request, 'service/index.html', context)

@login_required(login_url='/login/')
@all_user_required
def view_infos_specifique(request, projet_id, specific_id):
    projet = get_object_or_404(Projet, pk=projet_id)
    user = request.user
    # specific_infos = InfosSpecific.objects.get(utilisateur=user)
    specific_infos = get_object_or_404(InfosSpecific, pk=specific_id)
    partenaires_test = Partenaire.objects.filter(id_infos_specifique=specific_id)
    objectifs_test = Objectif.objects.filter(id_infos_specifique=specific_id)
    rez_test = Resultat.objects.filter(id_specific=specific_id)

    context = {
        'user': user,
        'projet': projet,
        'specific_infos': specific_infos,
        'partenaires_test': partenaires_test,
        'objectifs_test': objectifs_test,
        'rez_test': rez_test
    }
    return render(request, 'service/view_infos_specifique.html', context)

@login_required(login_url='/login/')
@all_user_required
def view_situation(request, projet_id, situation_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    situation = get_object_or_404(Situation, pk=situation_id)
    audit = Audit.objects.filter(situation=situation_id)
    context = {
        'user': user,
        'projet': projet,
        'situation': situation,
        'audit': audit
    }
    return render(request, 'service/view_situation.html', context)

@login_required(login_url='/login/')
@all_user_required
def view_planification(request, projet_id, activite_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    activite = get_object_or_404(Activite, pk=activite_id)
    partenaires = Partenaire.objects.filter(id_activite=activite_id) 
    context = {
        'activite': activite,
        'projet': projet,
        'partenaires': partenaires,
        'user': user
    }

    return render(request, 'service/view_planification.html', context)

@login_required(login_url='/login/')
@all_user_required
def view_realisation(request, projet_id, activite_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    activite = get_object_or_404(Realisation, pk=activite_id)
    partenaires = Partenaire.objects.filter(id_realisation=activite_id) 
    context = {
        'activite': activite,
        'projet': projet,
        'partenaires': partenaires,
        'user': user
    }

    return render(request, 'service/view_realisation.html', context)

@login_required(login_url='/login/')
@all_user_required
def view_generale(request):
    user = request.user
    infos = InfosGenerale.objects.get(utilisateur=user)
    objectifs = Objectif.objects.filter(id_general=infos)
    partenariats = Partenariat.objects.filter(id_general=infos)
    context = {
        'infos': infos,
        'objectifs': objectifs,
        'user': user,
        'partenariats': partenariats
    }

    return render(request, 'service/view_generale.html', context)

@login_required(login_url='/login/')
@all_user_required
def profils(request):
    return render(request, 'service/profil.html')

@login_required(login_url='/login/')
@responsable_charger_projet_required
def ajouter_infos_specifique(request, projet_id):
    partenaires = ListPartenaire.objects.all()
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    titres = TitreActivite.objects.all()
    secteurs = Secteur.objects.all()
    domaines = []
    activites = []

    selected_secteur = request.POST.get('secteur')
    if selected_secteur:
        domaines = SousSecteur.objects.filter(secteur_id=selected_secteur)

    selected_domaine = request.POST.get('domaine')
    if selected_domaine:
        activites = TitreActivite.objects.filter(domaine=selected_domaine)

    context = {
        'selected_secteur': int(selected_secteur) if selected_secteur else None,
        'selected_domaine': int(selected_domaine) if selected_domaine else None,
        'domaines': domaines,
        'activites': activites,
        'projet': projet,
        'user': user,
        'secteurs': secteurs,
        'titres': titres,
        'partenaires': partenaires,
        'default_partenaire': ListPartenaire.objects.first()
    }
    
    if request.method == 'POST':
        new_activite = request.POST.get('autre_activite')
        secteur_id = request.POST.get('secteur')
        sous_secteur_id = request.POST.get('domaine')
        titre_nom_id = request.POST.get('activite')

        # Print statements for debugging
        print(f"secteur_id: {secteur_id}")
        print(f"sous_secteur_id: {sous_secteur_id}")
        print(f"titre_nom_id: {titre_nom_id}")

        # Ajout de vérifications pour les IDs
        try:
            secteur = get_object_or_404(Secteur, pk=secteur_id)
            sous_secteur = get_object_or_404(SousSecteur, pk=sous_secteur_id)
            if new_activite:
                new_this = TitreActivite.objects.create(titre=new_activite, domaine=sous_secteur)
                titre_activite = get_object_or_404(TitreActivite, pk=new_this.id)
            else:
                titre_activite = get_object_or_404(TitreActivite, pk=titre_nom_id)
        except Exception as e:
            messages.error(request, f"Erreur: {str(e)}")
            return render(request, 'service/ajouter_planification.html', context)

        titre_activites = titre_activite.titre

        depense_globale = request.POST.get('globale_depense')
        nom_projet = request.POST.get('projet_name')
        obj_principal = request.POST.get('objectif_principale')
        obj_secondaire = request.POST.getlist('objectifs[]')
        resultats = request.POST.getlist('resultats[]')

        date_debut = request.POST.get('date_debut')
        date_fin = request.POST.get('date_fin')
        cout = request.POST.get('cout')

        benef_homme = request.POST.get('nbre_benef_homme')
        benef_femme = request.POST.get('nbre_benef_femme')
        
        total_benef = request.POST.get('total_benef_hidden')
        partenaire = request.POST.get('partners')

        # verifications des prix
        couts = float(request.POST.get('cout', 0))
        depense = Depense.objects.filter(id_projet=projet).first()
        total_depenses = 0
        if depense:
            total_depenses = (
                depense.consommable_divers + 
                depense.salaire_avantages + 
                depense.equipement_materiel
            )
            autres_depenses = AutreDepense.objects.filter(depense=depense).aggregate(
                total=Sum('prix')
            )['total'] or 0
            total_depenses += autres_depenses
        
        total_realisations = Realisation.objects.filter(
            id_projet=projet,
            realisation="Oui"
        ).aggregate(
            total=Sum('cout_realisation')
        )['total'] or 0

        # Si on a les deux totaux, vérifier par rapport au budget
        if total_depenses > 0 and total_realisations > 0:
            total_global = total_depenses + total_realisations
            message = ""
            if total_global < couts:
                difference = couts - total_global
                message = (
                    f"Le budget proposé est supérieur au total des dépenses et réalisations :\n\n"
                    f"Budget proposé : {couts:,.0f} FCFA\n"
                    f"Total dépenses : {total_depenses:,.0f} FCFA\n"
                    f"Total réalisations : {total_realisations:,.0f} FCFA\n"
                    f"Total global : {total_global:,.0f} FCFA\n\n"
                    f"Différence : {difference:,.0f} FCFA"
                )
            elif total_global > couts:
                difference = total_global - couts
                message = (
                    f"Le budget proposé est inférieur au total des dépenses et réalisations :\n\n"
                    f"Budget proposé : {couts:,.0f} FCFA\n"
                    f"Total dépenses : {total_depenses:,.0f} FCFA\n"
                    f"Total réalisations : {total_realisations:,.0f} FCFA\n"
                    f"Total global : {total_global:,.0f} FCFA\n\n"
                    f"Différence : {difference:,.0f} FCFA"
                )

            if message:
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'status': 'error',
                        'message': message
                    })
                else:
                    messages.error(request, message)
                    return render(request, 'service/ajouter_infos_specifique.html', context)

        # Gestion du partenaire
        if partenaire == 'Autre':
            partenaire = request.POST.get('otherPartner')
            if not ListPartenaire.objects.filter(nom=partenaire).exists():
                new_save = ListPartenaire.objects.create(nom=partenaire)
                new_save.save()
        else:
            partenaire = request.POST.get('partners')

        if all([nom_projet, obj_principal]):
            # Vérifier l'existence juste avant la création
            if InfosSpecific.objects.filter(id_projet=projet).exists():
                messages.info(request, 'INFORMATIONS SPECIFIQUES: ces donnees existent deja dans la base de donnees')
                return redirect('infos_specifique', projet_id=projet.id)

            specific = InfosSpecific.objects.create(
                utilisateur=user,
                id_secteur=secteur,
                id_sous_secteur=sous_secteur,
                id_titre_activites=titre_activite,
                id_projet=projet,
                nom=nom_projet,
                date_debut=date_debut,
                date_fin=date_fin,
                budget=cout,
                depense_globale=depense_globale,
                objectifs_principals=obj_principal,
                benef_direct_homme=benef_homme,
                benef_direct_femme=benef_femme,
                total_benef_direct=total_benef,
                partenaires=partenaire
            )
            specific.save()

            for obj in obj_secondaire:
                objectif = Objectif.objects.create(objectifs=obj, id_infos_specifique=specific)
                objectif.save()
            
            for rez in resultats:
                resultat = Resultat.objects.create(resultats=rez, id_specific=specific)
                resultat.save()
            
            messages.success(request, 'INFORMATIONS SPECIFIQUES: les informations ont ete creer avec succes')
            return redirect('infos_specifique', projet_id=projet.id)
        
        else:
            messages.info(request, 'INFORMATIONS SPECIFIQUES: veuillez renseigner toutes les informations')
            return render(request, 'service/ajouter_infos_specifique.html', context)

    return render(request, 'service/ajouter_infos_specifique.html', context)        

        

@login_required(login_url='/login/')
@responsable_charger_projet_required
def ajouter_planification(request, projet_id):
    #partenaire = [
        #"OMS", "LIGHT FOR THE WORLD", "Organisation Dupont pour le Developpement  et la Solidarité (ODDS)", "Caritas Autriche", "MANOS UNIDAS", "Association Burkinabè de Fundraising",
        #"Programme Alimentaire Mondial (PAM)", "ERIKS DEVELOPMENT Partner", "USAID", "SOLIDAR SUISSE", "Catholic relief service (CRS)", "ECOLE DE BREMEN", "FONDATION DOCTEUR ELVIRE ENGEL",
        #"FRATERNITE FRANCISCAINE", "Association Kontak", "Mrg DBERNADI PIER GIORGIO", "Caritas Insbruck", "Conférence épiscopale Italie", "Solid Aid", "ACDI VOCA", "Plan International", "Caritas Espagne",
        #"CordAid", "FHRAOC", "Afrika action", "Misereor/Ker", "Expertise France", "Norvegian Refugee Council", "Donateur Privé", "ADA", "MISSIO MUNCHEN", "FONAENF", "MISSIO AUSTRIA", "Junta CyL"
    #]
    partenaires = ListPartenaire.objects.all()
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    titres = TitreActivite.objects.all()
    secteurs = Secteur.objects.all()
    domaines = []
    activites = []

    selected_secteur = request.POST.get('secteur')
    if selected_secteur:
        domaines = SousSecteur.objects.filter(secteur_id=selected_secteur)

    selected_domaine = request.POST.get('domaine')
    if selected_domaine:
        activites = TitreActivite.objects.filter(domaine=selected_domaine)

    context = {
        'selected_secteur': int(selected_secteur) if selected_secteur else None,
        'selected_domaine': int(selected_domaine) if selected_domaine else None,
        'domaines': domaines,
        'activites': activites,
        'projet': projet,
        'user': user,
        'secteurs': secteurs,
        'titres': titres,
        'partenaires': partenaires,
        'default_partenaire': ListPartenaire.objects.first()
    }

    if request.method == 'POST':
        new_activite = request.POST.get('autre_activite')
        secteur_id = request.POST.get('secteur')
        sous_secteur_id = request.POST.get('domaine')
        titre_nom_id = request.POST.get('activite')

        # Print statements for debugging
        print(f"secteur_id: {secteur_id}")
        print(f"sous_secteur_id: {sous_secteur_id}")
        print(f"titre_nom_id: {titre_nom_id}")

        # Ajout de vérifications pour les IDs
        try:
            secteur = get_object_or_404(Secteur, pk=secteur_id)
            sous_secteur = get_object_or_404(SousSecteur, pk=sous_secteur_id)
            if new_activite:
                new_this = TitreActivite.objects.create(titre=new_activite, domaine=sous_secteur)
                titre_activite = get_object_or_404(TitreActivite, pk=new_this.id)
            else:
                titre_activite = get_object_or_404(TitreActivite, pk=titre_nom_id)
        except Exception as e:
            messages.error(request, f"Erreur: {str(e)}")
            return render(request, 'service/ajouter_planification.html', context)

        titre_activites = titre_activite.titre

        unite = request.POST.get('unite_hidden')
        quantite = request.POST.get('quantite')
        region = request.POST.get('region')
        province = request.POST.get('province')
        commune = request.POST.get('commune')
        cout_realisation = request.POST.get('cout_realisation')
        contrib_benf = request.POST.get('contrib_benef')
        contrib_part = request.POST.get('contrib_part')
        nbre_benef_homme = request.POST.get('nbre_benef_homme')
        nbre_benef_femme = request.POST.get('nbre_benef_femme')
        total_direct = request.POST.get('total_benef_hidden')
        paroisse = request.POST.get('paroisse')
        partenaire = request.POST.get('partners')
        # tester le partenaire
        if partenaire == 'Autre':
            partenaire = request.POST.get('otherPartner')
        else:
            partenaire = request.POST.get('partners')

        if not ListPartenaire.objects.filter(nom=partenaire).exists():
            nouveau_part = ListPartenaire.objects.create(nom=partenaire)
            nouveau_part.save()

        if Activite.objects.filter(id_projet=projet, titre=titre_activites).exists():
            messages.info(request, "PLANIFICATION OPERATIONNELLE : cette activitée existe déjà dans la base de données")
            return render(request, 'service/ajouter_planification.html', context)
        else:
            activite = Activite.objects.create(
                utilisateur=user,
                id_secteur=secteur,
                id_sous_secteur=sous_secteur,
                id_titre_activites = titre_activite,
                titre=titre_activites,
                unite_physique=unite,
                quantite_prevue=quantite,
                commune = commune,
                paroisse=paroisse,
                province=province,
                region=region,
                cout_realisation=cout_realisation,
                contribution_beneficiaire=contrib_benf,
                contribution_partenaire=contrib_part,
                nbre_benef_direct_homme=nbre_benef_homme,
                nbre_benef_direct_femme=nbre_benef_femme,
                total_benef_direct=total_direct,
                partenaires=partenaire,
                id_projet=projet,
                planification = "Oui"
            )

            realisation = Realisation.objects.create(
                id_projet=projet,
                utilisateur=user,
                id_secteur=secteur,
                id_sous_secteur=sous_secteur,
                id_titre_activites = titre_activite,
                titre=titre_activites,
                unite_physique=unite,
                commune = commune,
                paroisse=paroisse,
                province=province,
                region=region,
                cout_realisation=cout_realisation,
                contribution_beneficiaire=contrib_benf,
                contribution_partenaire=contrib_part,
                nbre_benef_direct_homme=nbre_benef_homme,
                nbre_benef_direct_femme=nbre_benef_femme,
                total_benef_direct=total_direct,
                nbre_benef_jeune_homme = 0,
                nbre_benef_jeune_femme = 0,
                total_benef_jeune = 0,
                partenaires=partenaire,
                realisation="Non",
                responsable='',
                quantite_prevue=0,
                id_activite=activite
            )
            realisation.save()
                
            messages.success(request, "PLANIFICATION OPERATIONNELLE : enregistrer avec succès")
            return redirect('planification', projet_id=projet.id)
    else:
        return render(request, 'service/ajouter_planification.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def ajouter_situation(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)

    context = {
        'user': user,
        'projet': projet
    }
    if request.method == 'POST':
        impot = request.POST.get('impot')
        cotisation = request.POST.get('cotisation')
        autre_contribution = request.POST.get('contribution')
        total = request.POST.get('total_hidden')

        designation = request.POST.getlist('designation[]')
        date_realisation = request.POST.getlist('date_realisation[]')
        nom_cabinet = request.POST.getlist('cabinet[]')

        if Situation.objects.filter(id_projet=projet).exists():
            messages.info(request, 'SITUATION FISCALE : existe déjà dans la base de données')
            return redirect('situation', projet_id=projet.id)
        else:
            situation = Situation.objects.create(
                utilisateur = user,
                id_projet = projet,
                titre = 'SITUATIONS FISCALES ET SOCIALES DU SED AU 31 DECEMBRE',
                impot = impot,
                cotisation = cotisation,
                autre_contribution = autre_contribution,
                total = total
            )
            for design, date, nom in zip(designation, date_realisation, nom_cabinet):
                situ = Audit.objects.create(situation=situation, designation=design, date_realisation=date, nom_cabinet=nom, titre_test='Derniers Audits comptables réalisés au cours de l\'année')
                situ.save()

            messages.success(request, 'SITUATION FISCALE : enregistrer avec succès')
            return redirect('situation', projet_id=projet.id)
    else:
        return render(request, 'service/ajouter_situation.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def ajouter_general(request):
    user = request.user

    #titres = TitreActivite.objects.all()
    context = {
            'user': user
        }
    if request.method == 'POST':
        nom_org = request.POST.get('nom_org')
        nature_org = request.POST.get('nature_org')
        sigle = request.POST.get('sigle')
        #pays_origine = request.POST.get('pays_ori')
        region = request.POST.get('region')
        province = request.POST.get('province')
        commune = request.POST.get('commune')
        village = request.POST.get('village')
        boite_postale = request.POST.get('boite_postale')
        numb_mobile = request.POST.get('ong_mobile')
        numb_fixe = request.POST.get('ong_fixe')
        adresse_mail = request.POST.get('ong_email_pro')
        site_web = request.POST.get('site_web')
        #
        nom_complet_resp = request.POST.get('resp_nom')
        nationalite_resp = request.POST.get('resp_nationalite')
        fonction_resp = request.POST.get('resp_fonction')
        numb_fixe_resp = request.POST.get('resp_fixe')
        numb_mobile_resp = request.POST.get('resp_mobile')
        #
        renou_instance = request.POST.get('date1')
        assem_general = request.POST.get('date2')
        session_statut = request.POST.get('date3')
        mandat_bureau = request.POST.get('date4')
        #
        nom_complet_canevas = request.POST.get('repondant_nom')
        numb_fixe_canevas = request.POST.get('repondant_fixe')
        numb_mobile_canevas = request.POST.get('repondant_mobile')
        adresse_mail_canevas = request.POST.get('repondant_adresse')
        #
        objectifs = request.POST.getlist('objectifs[]')
        #
        groupes_cibles = request.POST.getlist('groupe_check[]')
        groupes_test = ', '.join(groupes_cibles)
        autre_groupe = request.POST.get('autre_groupe')
        #
        total_pers_homme = request.POST.get('total_pers_homme')
        total_pers_femme = request.POST.get('total_pers_femme')
        #
        em_nation_cdi_homme = request.POST.get('em_nation_cdi_homme')
        em_nation_cdi_femme = request.POST.get('em_nation_cdi_femme')
        #
        em_nation_cdd_homme = request.POST.get('em_nation_cdd_homme')
        em_nation_cdd_femme = request.POST.get('em_nation_cdd_femme')
        #
        em_expa_cdi_homme = request.POST.get('em_expa_cdi_homme')
        em_expa_cdi_femme = request.POST.get('em_expa_cdi_femme')
        #
        em_expa_cdd_homme = request.POST.get('em_expa_cdd_homme')
        em_expa_cdd_femme = request.POST.get('em_expa_cdd_femme')
        #
        benevol_nation_homme = request.POST.get('benevol_nation_homme')
        benevol_nation_femme = request.POST.get('benevol_nation_femme')
        #
        benevol_expa_homme = request.POST.get('benevol_expa_homme')
        benevol_expa_femme = request.POST.get('benevol_expa_femme')
        #
        personnel_admin_homme = request.POST.get('personnel_admin_homme')
        personnel_admin_femme = request.POST.get('personnel_admin_femme')
        # ajouter un ou plusieurs partenaires
        nom_partenariat = request.POST.getlist('partenariat[]')
        numero_partenariat = request.POST.getlist('number_part[]')
        date_debut_part = request.POST.getlist('date_debut_part[]')
        date_fin_part = request.POST.getlist('date_fin_part[]')

        if InfosGenerale.objects.filter(utilisateur=user).exists():
            messages.info(request, 'INFORMATIONS GENERALES : existe déjà dans la base de données')
            return redirect('generales')
        
        else:
            if all([ nom_org, nature_org ]):
                infos = InfosGenerale.objects.create(
                    utilisateur = user,
                    nom_org = nom_org,
                    nature_org = nature_org,
                    sigle = sigle,
                    pays_origine = 'Burkina-Faso',
                    region = region,
                    province = province,
                    commune = commune,
                    village = village,
                    boite_postale = boite_postale,
                    numb_mobile = numb_mobile,
                    numb_fixe = numb_fixe,
                    adresse_mail = adresse_mail,
                    site_web = site_web,
                    nom_complet_resp = nom_complet_resp,
                    nationalite_resp = nationalite_resp,
                    fonction_resp = fonction_resp,
                    numb_fixe_resp = numb_fixe_resp,
                    numb_mobile_resp = numb_mobile_resp,
                    renou_instance = renou_instance,
                    assem_general = assem_general,
                    session_statut = session_statut,
                    mandat_bureau = mandat_bureau,
                    nom_complet_canevas = nom_complet_canevas,
                    numb_fixe_canevas = numb_fixe_canevas,
                    numb_mobile_canevas = numb_mobile_canevas,
                    adresse_mail_canevas = adresse_mail_canevas,
                    groupes_cibles = groupes_test,
                    total_pers_homme = total_pers_homme,
                    total_pers_femme = total_pers_femme,
                    em_nation_cdi_homme = em_nation_cdi_homme,
                    em_nation_cdi_femme = em_nation_cdi_femme,
                    em_nation_cdd_homme = em_nation_cdd_homme,
                    em_nation_cdd_femme = em_nation_cdd_femme,
                    em_expa_cdi_homme = em_expa_cdi_homme,
                    em_expa_cdi_femme = em_expa_cdi_femme,
                    em_expa_cdd_homme = em_expa_cdd_homme,
                    em_expa_cdd_femme = em_expa_cdd_femme,
                    benevol_nation_homme = benevol_nation_homme,
                    benevol_nation_femme = benevol_nation_femme,
                    benevol_expa_homme = benevol_expa_homme,
                    benevol_expa_femme = benevol_expa_femme,
                    personnel_admin_homme = personnel_admin_homme,
                    personnel_admin_femme = personnel_admin_femme,
                    autre_groupe = autre_groupe
                )
                infos.save()
                for obj in objectifs:
                    objectif = Objectif.objects.create(
                        objectifs=obj, 
                        id_general=infos
                    )
                    objectif.save()

                for name, numero, date_debut, date_fin in zip(nom_partenariat, numero_partenariat, date_debut_part, date_fin_part):
                    partenariat = Partenariat.objects.create(
                        nom = name,
                        numero = numero,
                        date_debut = date_debut,
                        date_fin = date_fin,
                        id_general = infos
                    )
                    partenariat.save()

                messages.success(request, 'INFORMATIONS GENERALES : créer avec succès')
                return redirect('generales')
            
            else:
                messages.info(request, 'INFORMATIONS GENERALES : veuillez renseigner toutes les informations')
                return render(request, 'service/ajouter_general.html', context)
            
    return render(request, 'service/ajouter_general.html', context)


@login_required(login_url='/login/')
@responsable_charger_projet_required
def modifier_realisation(request, projet_id, activite_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    activite = get_object_or_404(Realisation, pk=activite_id)
    partenaires = ListPartenaire.objects.all()
    
    context = {
        'activite': activite,
        'projet': projet,
        'user': user,
        'partenaires': partenaires
        }
    if request.method == 'POST':
        quantite_realise = request.POST.get('quantite')
        date_debut = request.POST.get('periode_debut')
        date_fin = request.POST.get('periode_fin')
        responsable = request.POST.get('responsable')

        benef_homme_jeu = request.POST.get('nbre_ben_homme_jeu')
        benef_femme_jeu = request.POST.get('nbre_ben_femme_jeu')
        total_benef_jeu = request.POST.get('total_benef_jeu_hidden')

        commune = request.POST.get('commune')
        paroisse = request.POST.get('paroisse')
        cout_total = request.POST.get('cout_realisation')
        contrib_benef = request.POST.get('contrib_benef')
        contrib_part = request.POST.get('contrib_part')
        benef_direct_homme = request.POST.get('nbre_benef_homme')
        benef_direct_femme = request.POST.get('nbre_benef_femme')
        total_benef = request.POST.get('total_benef_hidden')
        partenaire = request.POST.get('partners')

        cout_total = float(request.POST.get('cout_realisation', 0))
        # Vérifier si on a un budget
        infos = InfosSpecific.objects.filter(id_projet=projet).first()

        # Vérifier si on a des dépenses
        depense = Depense.objects.filter(id_projet=projet).first()
        total_depenses = 0
        if depense:
            total_depenses = (
                depense.consommable_divers + 
                depense.salaire_avantages + 
                depense.equipement_materiel
            )
            autres_depenses = AutreDepense.objects.filter(depense=depense).aggregate(
                total=Sum('prix')
            )['total'] or 0
            total_depenses += autres_depenses

        # On ne vérifie que si on a le budget ET les dépenses
        if infos and infos.budget and total_depenses > 0:
            # Calculer le total des autres réalisations
            total_autres_realisations = Realisation.objects.filter(
                id_projet=projet,
                realisation="Oui"
            ).exclude(
                id=activite_id
            ).aggregate(
                total=Sum('cout_realisation')
            )['total'] or 0

            total_global = total_depenses + total_autres_realisations + cout_total
            
            if total_global != infos.budget:
                message = ""
                if total_global < infos.budget:
                    difference = infos.budget - total_global
                    message = (
                        f"Le total serait inférieur au budget :\n\n"
                        f"Budget : {infos.budget:,.0f} FCFA\n"
                        f"Dépenses : {total_depenses:,.0f} FCFA\n"
                        f"Autres réalisations : {total_autres_realisations:,.0f} FCFA\n"
                        f"Cette réalisation : {cout_total:,.0f} FCFA\n"
                        f"Total : {total_global:,.0f} FCFA\n\n"
                        f"Différence : {difference:,.0f} FCFA"
                    )
                else:
                    difference = total_global - infos.budget
                    message = (
                        f"Le total dépasserait le budget :\n\n"
                        f"Budget : {infos.budget:,.0f} FCFA\n"
                        f"Dépenses : {total_depenses:,.0f} FCFA\n"
                        f"Autres réalisations : {total_autres_realisations:,.0f} FCFA\n"
                        f"Cette réalisation : {cout_total:,.0f} FCFA\n"
                        f"Total : {total_global:,.0f} FCFA\n\n"
                        f"Dépassement : {difference:,.0f} FCFA"
                    )

                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'status': 'error',
                        'message': message
                    })
                else:
                    messages.error(request, message)
                    return render(request, 'service/modifier_realisation.html', context)
        #
        if partenaire == 'Autre':
            partenaire = request.POST.get('otherPartner')
            if not ListPartenaire.objects.filter(nom=partenaire).exists():
                new_save = ListPartenaire.objects.create(nom=partenaire)
                new_save.save()
        else:
            partenaire = request.POST.get('partners')

        if all([quantite_realise]):
            with transaction.atomic():
                activite = Realisation.objects.get(id=activite_id)
                activite.utilisateur = user
                activite.quantite_prevue = quantite_realise
                activite.periode_prevue_debut = date_debut
                activite.periode_prevue_fin = date_fin
                activite.responsable = responsable
                activite.id_projet = projet
                activite.commune = commune
                activite.paroisse = paroisse
                activite.cout_realisation = cout_total
                activite.contribution_beneficiaire = contrib_benef
                activite.contribution_partenaire = contrib_part
                activite.nbre_benef_direct_homme = benef_direct_homme
                activite.nbre_benef_direct_femme = benef_direct_femme
                activite.total_benef_direct = total_benef
                activite.nbre_benef_jeune_homme = benef_homme_jeu
                activite.nbre_benef_jeune_femme = benef_femme_jeu
                activite.total_benef_jeune = total_benef_jeu
                activite.partenaires = partenaire
                activite.realisation = "Oui"
                # sauvegarde
                activite.save()
                
            messages.success(request, "RÉALISATION : activité realisée avec succès")
            return redirect('realisation', projet_id=projet.id)
        else:
            messages.info(request, 'Erreur: Veuillez remplir tout les champs.')
            return render(request, 'service/modifier_realisation.html', context)
        
    return render(request, 'service/modifier_realisation.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def modifier_general(request):
    user = request.user
    infos = InfosGenerale.objects.get(utilisateur=user)
    objectifs = Objectif.objects.filter(id_general=infos)
    partenariats = Partenariat.objects.filter(id_general=infos)
    context = {
            'infos': infos,
            'objectifs': objectifs,
            'user': user,
            'partenariats': partenariats
        }
    
    if request.method == 'POST':
        nom_org = request.POST.get('nom_org')
        nature_org = request.POST.get('nature_org')
        sigle = request.POST.get('sigle')
        #pays_origine = request.POST.get('pays_ori')
        region = request.POST.get('region')
        province = request.POST.get('province')
        commune = request.POST.get('commune')
        village = request.POST.get('village')
        boite_postale = request.POST.get('boite_postale')
        numb_mobile = request.POST.get('ong_mobile')
        numb_fixe = request.POST.get('ong_fixe')
        adresse_mail = request.POST.get('ong_email_pro')
        site_web = request.POST.get('site_web')
        #
        nom_complet_resp = request.POST.get('resp_nom')
        nationalite_resp = request.POST.get('resp_nationalite')
        fonction_resp = request.POST.get('resp_fonction')
        numb_fixe_resp = request.POST.get('resp_fixe')
        numb_mobile_resp = request.POST.get('resp_mobile')
        #
        renou_instance = request.POST.get('date1')
        assem_general = request.POST.get('date2')
        session_statut = request.POST.get('date3')
        mandat_bureau = request.POST.get('date4')
        #
        nom_complet_canevas = request.POST.get('repondant_nom')
        numb_fixe_canevas = request.POST.get('repondant_fixe')
        numb_mobile_canevas = request.POST.get('repondant_mobile')
        adresse_mail_canevas = request.POST.get('repondant_adresse')
        #
        objectifs = request.POST.getlist('objectifs[]')
        #
        groupes_cibles = request.POST.getlist('groupe_check[]')
        #
        total_pers_homme = request.POST.get('total_pers_homme')
        total_pers_femme = request.POST.get('total_pers_femme')
        #
        em_nation_cdi_homme = request.POST.get('em_nation_cdi_homme')
        em_nation_cdi_femme = request.POST.get('em_nation_cdi_femme')
        #
        em_nation_cdd_homme = request.POST.get('em_nation_cdd_homme')
        em_nation_cdd_femme = request.POST.get('em_nation_cdd_femme')
        #
        em_expa_cdi_homme = request.POST.get('em_expa_cdi_homme')
        em_expa_cdi_femme = request.POST.get('em_expa_cdi_femme')
        #
        em_expa_cdd_homme = request.POST.get('em_expa_cdd_homme')
        em_expa_cdd_femme = request.POST.get('em_expa_cdd_femme')
        #
        benevol_nation_homme = request.POST.get('benevol_nation_homme')
        benevol_nation_femme = request.POST.get('benevol_nation_femme')
        #
        benevol_expa_homme = request.POST.get('benevol_expa_homme')
        benevol_expa_femme = request.POST.get('benevol_expa_femme')
        #
        personnel_admin_homme = request.POST.get('personnel_admin_homme')
        personnel_admin_femme = request.POST.get('personnel_admin_femme')
        # ajouter un ou plusieurs partenaires
        nom_partenariat = request.POST.getlist('partenariat[]')
        numero_partenariat = request.POST.getlist('number_part[]')
        date_debut_part = request.POST.getlist('date_debut_part[]')
        date_fin_part = request.POST.getlist('date_fin_part[]')
        

        if all([
                nom_org,
                nature_org,
                sigle,
                region,
                province,
                commune,
                village,
                boite_postale,
                numb_fixe,
                numb_mobile,
                adresse_mail,
                site_web,
                nom_complet_resp,
                nationalite_resp,
                fonction_resp,
                numb_fixe_resp,
                numb_mobile_resp,
                renou_instance,
                assem_general,
                session_statut,
                mandat_bureau,
                nom_complet_canevas,
                numb_fixe_canevas,
                numb_mobile_canevas,
                adresse_mail_canevas,
                groupes_cibles,
                total_pers_homme,
                total_pers_femme,
                em_nation_cdi_homme,
                em_nation_cdi_femme,
                em_nation_cdd_homme,
                em_nation_cdd_femme,
                em_expa_cdi_homme,
                em_expa_cdi_femme,
                em_expa_cdd_homme,
                em_expa_cdd_femme,
                benevol_nation_homme,
                benevol_nation_femme,
                benevol_expa_homme,
                benevol_expa_femme,
                personnel_admin_homme,
                personnel_admin_femme
            ]):
            with transaction.atomic():
                generales = InfosGenerale.objects.get(utilisateur=user)
                # mettre a jour les champs
                generales.nom_org = nom_org
                generales.nature_org = nature_org
                generales.sigle = sigle
                generales.region = region
                generales.province = province
                generales.commune = commune
                generales.village = village
                generales.boite_postale = boite_postale
                generales.numb_mobile = numb_mobile
                generales.numb_fixe = numb_fixe
                generales.adresse_mail = adresse_mail
                generales.site_web = site_web
                generales.nom_complet_resp = nom_complet_resp
                generales.nationalite_resp = nationalite_resp
                generales.fonction_resp = fonction_resp
                generales.numb_fixe_resp = numb_fixe_resp
                generales.numb_mobile_resp = numb_mobile_resp
                generales.renou_instance = renou_instance
                generales.assem_general = assem_general
                generales.session_statut = session_statut
                generales.mandat_bureau = mandat_bureau
                generales.nom_complet_canevas = nom_complet_canevas
                generales.numb_fixe_canevas = numb_fixe_canevas
                generales.numb_mobile_canevas = numb_mobile_canevas
                generales.adresse_mail_canevas = adresse_mail_canevas
                generales.groupes_cibles = groupes_cibles
                generales.total_pers_homme = total_pers_homme
                generales.total_pers_femme = total_pers_femme
                generales.em_nation_cdi_homme = em_nation_cdi_homme
                generales.em_nation_cdi_femme = em_nation_cdi_femme
                generales.em_nation_cdd_homme = em_nation_cdd_homme
                generales.em_nation_cdd_femme = em_nation_cdd_femme
                generales.em_expa_cdi_homme = em_expa_cdi_homme
                generales.em_expa_cdi_femme = em_expa_cdi_femme
                generales.em_expa_cdd_homme = em_expa_cdd_homme
                generales.em_expa_cdd_femme = em_expa_cdd_femme
                generales.benevol_nation_homme = benevol_nation_homme
                generales.benevol_nation_femme = benevol_nation_femme
                generales.benevol_expa_homme = benevol_expa_homme
                generales.benevol_expa_femme = benevol_expa_femme
                generales.personnel_admin_homme = personnel_admin_homme
                generales.personnel_admin_femme = personnel_admin_femme
                # sauvegarder
                generales.save()

                # objectifs
                for obj in objectifs:
                    try:
                        Objectif.objects.get(objectifs=obj, id_general=generales)
                    except ObjectDoesNotExist:
                        Objectif.objects.create(objectifs=obj, id_general=generales)
                
                # partenariats
                for name, numero, date_debut, date_fin in zip(nom_partenariat, numero_partenariat, date_debut_part, date_fin_part):
                    try:
                        Partenariat.objects.get(
                            nom = name,
                            numero = numero,
                            date_debut = date_debut,
                            date_fin = date_fin,
                            id_general = generales
                        )
                    except ObjectDoesNotExist:
                        partenariat = Partenariat.objects.create(
                            nom = name,
                            numero = numero,
                            date_debut = date_debut,
                            date_fin = date_fin,
                            id_general = generales
                        )

        messages.success(request, 'INFORMATIONS GENERALES : modifier avec succès')
        #return render(request, 'service/choix_projet.html', context)
        return redirect('generales')

    else:
        return render(request, 'service/modifier_general.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def modifier_infos_specifique(request, projet_id, specific_id):

    partenaires = ListPartenaire.objects.all()

    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    specific_infos = get_object_or_404(InfosSpecific, pk=specific_id)
    objectifs_test = Objectif.objects.filter(id_infos_specifique=specific_id)
    rez_test = Resultat.objects.filter(id_specific=specific_id)

    context = {
        'user': user,
        'projet': projet,
        'specific_infos': specific_infos,
        'objectifs_test': objectifs_test,
        'rez_test': rez_test,
        'partenaires': partenaires
    }

    if request.method == 'POST':
        nom_projet = request.POST.get('projet_name')
        obj_principal = request.POST.get('objectif_principale')
        obj_secondaire = request.POST.getlist('objectifs[]')
        resultats = request.POST.getlist('resultats[]')

        date_debut = request.POST.get('date_debut')
        date_fin = request.POST.get('date_fin')
        budget = request.POST.get('cout')
        depense_globale = request.POST.get('globale_depense')

        benef_homme = request.POST.get('nbre_ben_homme')
        benef_femme = request.POST.get('nbre_ben_femme')
        partenaire = request.POST.get('partners')
        #
        if partenaire == 'Autre':
            partenaire = request.POST.get('otherPartner')
            if not ListPartenaire.objects.filter(nom=partenaire).exists():
                new_save = ListPartenaire.objects.create(nom=partenaire)
                new_save.save()
        else:
            partenaire = request.POST.get('partners')

        #nom_partenaire = request.POST.getlist('nom_partenaire[]')
        #part_partenaire = request.POST.getlist('part_partenaire[]')

        if all([nom_projet, obj_principal, date_debut, date_fin, budget]):
            with transaction.atomic():
                specific = InfosSpecific.objects.get(id=specific_id)
                specific.utilisateur = user
                specific.nom = nom_projet
                specific.objectifs_principals = obj_principal
                specific.date_debut = date_debut
                specific.date_fin = date_fin
                specific.budget = budget
                specific.depense_globale = depense_globale
                specific.benef_direct_homme = benef_homme
                specific.benef_direct_femme = benef_femme
                specific.partenaires = partenaire
                specific.save()

                for obj in obj_secondaire:
                    try:
                        Objectif.objects.get(objectifs=obj, id_infos_specifique=specific)
                    except ObjectDoesNotExist:
                        objectif = Objectif.objects.create(objectifs=obj, id_infos_specifique=specific)
                        objectif.save()
                
                for rez in resultats:
                    try:
                        Resultat.objects.get(resultats=rez, id_specific=specific)
                    except ObjectDoesNotExist:
                        resultat = Resultat.objects.create(resultats=rez, id_specific=specific)
                        resultat.save()
                
        messages.success(request, 'INFORMATIONS SPECIFIQUES : modifier avec succès')
        return redirect('infos_specifique', projet_id=projet.id)
        
    else:
        return render(request, 'service/modifier_infos_specifique.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def modifier_situation(request, projet_id, situation_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    situation = get_object_or_404(Situation, pk=situation_id)
    audits = Audit.objects.filter(situation=situation_id)
    
    context = {
        'user': user,
        'projet': projet,
        'situation': situation,
        'audits': audits
    }
    if request.method == 'POST':
        impot_test = request.POST.get('impot')
        cotisation_test = request.POST.get('cotisation')
        contribution_test = request.POST.get('contribution')

        total = request.POST.get('total_hidden')

        designation_test = request.POST.getlist('designation[]')
        date_realisation_test = request.POST.getlist('date_realisation[]')
        nom_cabinet_test = request.POST.getlist('cabinet[]')

        if all([
            impot_test, cotisation_test, contribution_test,
            designation_test, date_realisation_test, nom_cabinet_test
        ]):
            with transaction.atomic():
                #situt = Situation.objects.get(id=situation_id)
                situation.impot = impot_test
                situation.cotisation = cotisation_test
                situation.autre_contribution = contribution_test
                situation.total = total
                situation.save()

                for design, date, nom in zip(designation_test, date_realisation_test, nom_cabinet_test):
                    try:
                        Audit.objects.get(
                            situation = situation, 
                            designation = design, 
                            date_realisation = date, 
                            nom_cabinet = nom
                        )
                        messages.info(request, 'SITUATION FISCALE : existe deja dans la base de données')
                    except ObjectDoesNotExist:
                        auditer = Audit.objects.create(
                            situation = situation,
                            designation = design,
                            date_realisation = date,
                            nom_cabinet = nom,
                            titre_test='Derniers Audits comptables réalisés au cours de l\'année'
                        )
                        auditer.save()

        messages.success(request, 'SITUATION FISCALE : modifier avec succès')
        return render(request, 'service/choix_projet.html', context)
    else:
        return render(request, 'service/modifier_situation.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def modifier_planification(request, projet_id, activite_id):
    
    partenaires = ListPartenaire.objects.all()
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    activite = get_object_or_404(Activite, pk=activite_id)
    context = {
        'activite': activite,
        'projet': projet,
        'partenaires': partenaires,
        'user': user,
        # 'unite_physique': unite_physique,
        'partenaires': partenaires
    }
    if request.method == 'POST':
        unite = request.POST.get('unite')
        quantite = request.POST.get('quantite')
        region = request.POST.get('region')
        province = request.POST.get('province')
        commune = request.POST.get('commune')
        cout_realisation = request.POST.get('cout_realisation')
        contrib_benef = request.POST.get('contrib_benef')
        contrib_part = request.POST.get('contrib_part')
        total_direct = request.POST.get('total_benef')
        paroisse = request.POST.get('paroisse')
        partenaire = request.POST.get('partners')
        # test
        if partenaire == 'Autre':
            partenaire = request.POST.get('otherPartner')
            if not ListPartenaire.objects.filter(nom=partenaire).exists():
                new_save = ListPartenaire.objects.create(nom=partenaire)
                new_save.save()
        else:
            partenaire = request.POST.get('partners')


        if all([unite, quantite, region, partenaire]):
            with transaction.atomic():
                activite = Activite.objects.get(id=activite_id)
                # Mettre a jours les champs de l'activite
                activite.unite_physique = unite
                activite.quantite_prevue = quantite
                activite.region = region
                activite.province = province
                activite.commune = commune
                activite.cout_realisation = cout_realisation
                activite.contribution_beneficiaire = contrib_benef
                activite.paroisse = paroisse
                activite.contribution_partenaire = contrib_part
                activite.total_benef_direct = total_direct
                activite.partenaires = partenaire
                # sauvegarder les modifications
                activite.save()
            
        messages.success(request, 'PLANIFICATION OPERATIONNELLE : modifier avec succès')
        return render(request, 'service/choix_projet.html', context)
    else:
        return render(request, 'service/modifier_planification.html', context)


@login_required(login_url='/login/')
@responsable_charger_projet_required
def create_projet(request):
    user = request.user
    context = {
        'user': user
    }
    if request.method == 'POST':
        total_projet = Projet.objects.filter(utilisateur=user).count()
        new_numb = total_projet + 1
        new_name = f"Projet {new_numb}"

        project = Projet.objects.create(nom=new_name, utilisateur=user)

        messages.success(request, 'PROJET : créer avec succès')
        return redirect('choix_projet', projet_id=project.id)
    else:
        return render(request, 'service/create_projet.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def choisir_ajouter(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)

    context = {
        'projet': projet,
        'user': user
    }
    return render(request, 'service/choisir_ajouter.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def choisir_modifier(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)

    context = {
        'projet': projet,
        'user': user
    }
    return render(request, 'service/choisir_modifier.html', context)

@login_required(login_url='/login/')
@all_user_required
def projet(request):
    user = request.user
    #projects = Projet.objects.filter(utilisateur=user)
    projects = Projet.objects.all()
    #project_users = Projet.objects.filter(utilisateur=user)
    context = {
        'projects': projects,
        'user': user
    }
    return render(request, 'service/projet.html', context)

@login_required(login_url='/login/')
@all_user_required
def infos_specifique(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    specifique = InfosSpecific.objects.filter(id_projet=projet.id)
    context = {
        'user': user,
        'projet': projet,
        'specifique': specifique
    }
    return render(request, 'service/infos_specifique.html', context)

@login_required(login_url='/login/')
@all_user_required
def situation(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    #situations = Situation.objects.filter(utilisateur=user)
    situations = Situation.objects.filter(id_projet=projet.id)
    context = {
        'user': user,
        'projet': projet,
        'situations': situations
    }
    return render(request, 'service/situation.html', context)

@login_required(login_url='/login/')
@all_user_required
def activites(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    # filtrer les activites suivant les users
    #activites = Activite.objects.filter(utilisateur=user)
    activites = Activite.objects.filter(id_projet=projet.id)

    context = {
        'projet': projet,
        'activites': activites,
        'user': user
    }
    return render(request, 'service/planification.html', context)

@login_required(login_url='/login/')
@all_user_required
def realisation(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    #activites = Realisation.objects.filter(utilisateur=user)
    activites = Realisation.objects.filter(id_projet=projet.id)

    context = {
        'projet': projet,
        'activites': activites,
        'user': user
    }
    return render(request, 'service/realisation.html', context)

@login_required(login_url='/login/')
@all_user_required
def choix_projet(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    #activite = Activite.objects.all()
    activite = Activite.objects.filter(id_projet=projet.id)

    context = {
        'projet': projet,
        'activite': activite,
        'user': user
    }

    return render(request, 'service/choix_projet.html', context)

@login_required(login_url='/login/')
@all_user_required
def generales(request):
    user = request.user
    projet = Projet.objects.filter(utilisateur=user)
    #projet = Projet.objects.all()
    infos_generale = InfosGenerale.objects.all()
    has_infos_general = infos_generale.exists()
    #infos_generale = InfosGenerale.objects.filter(id_projet=projet.id)
    #infos_generale_users = InfosGenerale.objects.filter(utilisateur=user)

    context = {
        'projet': projet,
        'infos_generale': infos_generale,
        'user': user,
        'has_infos_general': has_infos_general
    }

    return render(request, 'service/generales.html', context)

@login_required(login_url='/login/')
@all_user_required
def change_password(request):
    if request.method == 'POST':
        old_pass = request.POST.get('old_password')
        new_pass1 = request.POST.get('new_password1')
        new_pass2 = request.POST.get('new_password2')

        #  verifier l'ancien mdp
        if not request.user.check_password(old_pass):
            messages.info(request, 'Votre ancien mot de passe est incorrect')
            return redirect('profils')
        
        # verifier les deux nouveaux mdp
        if new_pass1 != new_pass2:
            messages.info(request, 'Les nouveaux mots de passe ne correspondent pas')
            return redirect('profils')
        
        request.user.set_password(new_pass1)
        request.user.save()
        update_session_auth_hash(request, request.user)
        messages.success(request, 'Votre mot de passe a été mis à jour avec succès.')
        return redirect('profils')
        
    return  render(request, 'service/profil.html')

@login_required(login_url='/login/')
@responsable_charger_projet_required
def get_domaines(request):
    secteur_id = request.GET.get('secteur_id')
    domaines = SousSecteur.objects.filter(secteur_id=secteur_id).values('id', 'titre')
    return JsonResponse(list(domaines), safe=False)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def get_activites(request):
    domaine_id = request.GET.get('domaine_id')
    activites = TitreActivite.objects.filter(domaine_id=domaine_id).values('id', 'titre', 'unite_physique')
    return JsonResponse(list(activites), safe=False)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def get_activite_details(request):
    activite_id = request.GET.get('id')
    activite = get_object_or_404(Activite, pk=activite_id)
    # titre = Activite.objects.filter(titre=activite.titre)
    try:
        activite_te = activite.titre
        activite_test = Activite.objects.filter(titre=activite_te)
        data = {
            
            # 'unite_physique': activite_test.unite_physique,
            # 'quantite_prevue': activite_test.quantite_prevue,
            # 'periode_prevue_debut': activite_test.periode_prevue_debut,
            # 'periode_prevue_fin': activite_test.periode_prevue_fin,
            # 'responsable': activite_test.responsable,
            # 'part_burkina': activite_test.part_burkina
        }
        return JsonResponse(data)
    except Activite.DoesNotExist:
        return JsonResponse({'error': 'Activite non trouver'}, status=404)
    

@login_required(login_url='/login/')
@responsable_charger_projet_required
def get_activite(request):
    user = request.user
    activite_titre = request.GET.get('titre')
    projet_id = request.GET.get('projet_id')
    project = get_object_or_404(Projet, pk=projet_id)
    if activite_titre:
        # activite = Activite.objects.get(titre=activite_titre, utilisateur=user, id_projet=project)
        activite = TitreActivite.objects.get(titre=activite_titre)

        unite_physique = activite.unite_physique
        data = {
                'titre': activite.titre,
                'unite_physique': unite_physique
        }
        return JsonResponse(data)
        
    return JsonResponse({'error': 'Titre non fourni'}, status=400)


@login_required(login_url='/login/')
@all_user_required
def documentation(request):
    return render(request, 'service/documentation.html')

@login_required(login_url='/login/')
@all_user_required
def return_pdf(request):
    pdf_path = os.path.join(settings.BASE_DIR, 'services', 'files', 'Documentation.pdf')
    
    # Ouvre le fichier PDF
    with open(pdf_path, 'rb') as pdf_file:
        # Crée une réponse HTTP
        response = HttpResponse(pdf_file.read(), content_type='application/pdf')
        response['Content-Disposition'] = 'inline; filename="Documentation.pdf"'
        
        return response
    

@login_required(login_url='/login/')
@responsable_required
def utilisateur(request):
    users = User.objects.exclude(is_superuser=True)
    context = {
        'users': users
    }
    return render(request, 'service/utilisateur.html', context)

@login_required(login_url='/login/')
@responsable_required
def modifier_utilisateur(request):
    if request.method == 'POST':
        user_id = request.POST.get('user_id')
        user = get_object_or_404(User, pk=user_id)
        type_name = request.POST.get('user_type')

        user.is_user = False
        user.is_gestion = False
        user.is_responsable = False
        user.is_assistant = False
        user.is_charger_programme = False
        user.is_charger_projet = False

        if not type_name:
            messages.info(request, "Veuillez selectionner un type d'utilisateur svp !")
            return redirect('utilisateur')
        elif type_name:
            if type_name == 'gestion':
                user.is_gestion = True
            elif type_name == 'responsable':
                user.is_responsable = True
            elif type_name == 'assistant':
                user.is_assistant = True
            elif type_name == 'programme':
                user.is_charger_programme = True
            elif type_name == 'projet':
                user.is_charger_projet = True

            user.save()
            messages.success(request, "Le type de l'utilisateur a été changer avec succès")
            return redirect('utilisateur')


@login_required(login_url='/login/')
@all_user_required
def depense(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    #depenses = Depense.objects.filter(utilisateur=user)
    depenses = Depense.objects.filter(id_projet=projet.id)
    nombre = depenses.count()
    context = {
        'user': user,
        'projet': projet,
        'depenses': depenses,
        'nombre': nombre
    }
    return render(request, 'service/depense.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def ajouter_depense(request, projet_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    context = {
        'user': user,
        'projet': projet
    }
    if request.method == 'POST':
        salaire = request.POST.get('salaire')
        consommable_divers = request.POST.get('consommable_diver')
        equipement_materiel = request.POST.get('equipement_materiel')
        
        intitules = request.POST.getlist('intitules[]')
        prix = request.POST.getlist('couts[]')

        salaires = float(request.POST.get('salaire', 0))
        consommable_diverss = float(request.POST.get('consommable_diver', 0))
        equipement_materiels = float(request.POST.get('equipement_materiel', 0))
        
        intituless = request.POST.getlist('intitules[]')
        prixs = request.POST.getlist('couts[]')
        total_autres = sum(float(p) for p in prixs if p)

        nouveau_total_depenses = salaires + consommable_diverss + equipement_materiels + total_autres
         # Vérifier si on a un budget
        infos = InfosSpecific.objects.filter(id_projet=projet).first()

         # Récupérer les réalisations si elles existent
        total_realisations = Realisation.objects.filter(
            id_projet=projet,
            realisation="Oui"
        ).aggregate(
            total=Sum('cout_realisation')
        )['total'] or 0

        # Si on a le budget et des réalisations
        if infos and infos.budget and total_realisations > 0:
            total_global = nouveau_total_depenses + total_realisations
            
            if total_global != infos.budget:
                message = ""
                if total_global < infos.budget:
                    difference = infos.budget - total_global
                    message = (
                        f"Le total est inférieur au budget :\n\n"
                        f"Budget : {infos.budget:,.0f} FCFA\n"
                        f"Dépenses proposées : {nouveau_total_depenses:,.0f} FCFA\n"
                        f"Réalisations existantes : {total_realisations:,.0f} FCFA\n"
                        f"Total : {total_global:,.0f} FCFA\n\n"
                        f"Différence : {difference:,.0f} FCFA"
                    )
                else:
                    difference = total_global - infos.budget
                    message = (
                        f"Le total dépasse le budget :\n\n"
                        f"Budget : {infos.budget:,.0f} FCFA\n"
                        f"Dépenses proposées : {nouveau_total_depenses:,.0f} FCFA\n"
                        f"Réalisations existantes : {total_realisations:,.0f} FCFA\n"
                        f"Total : {total_global:,.0f} FCFA\n\n"
                        f"Dépassement : {difference:,.0f} FCFA"
                    )

                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'status': 'error',
                        'message': message
                    })
                else:
                    messages.error(request, message)
                    return render(request, 'service/ajouter_depense.html', context)

        if Depense.objects.filter(id_projet=projet).exists():
            messages.info(request, "DEPENSES DE FONCTIONNEMENT : existe deja dans la base de donnees")
            return redirect('depense', projet_id=projet.id)
        else:
            depenses = Depense.objects.create(
                utilisateur = user,
                id_projet = projet,
                name='DEPENSES DE FONCTIONNEMENT',
                consommable_divers = consommable_divers,
                salaire_avantages = salaire,
                equipement_materiel = equipement_materiel  
            )
            depenses.save()
        
        if all([intitules, prix]):
            for nom, price in zip(intitules, prix):
                autres_depenses = AutreDepense.objects.create(
                    depense = depenses,
                    intitule = nom,
                    prix = price
                )
                autres_depenses.save()

            messages.success(request, "DEPENSES DE FONCTIONNEMENT : enregistrer avec succes")
            return redirect('depense', projet_id=projet.id)

    return render(request, 'service/ajouter_depense.html', context)

@login_required(login_url='/login/')
@responsable_charger_projet_required
def modifier_depense(request, projet_id, depense_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    depense = get_object_or_404(Depense, pk=depense_id)
    autre_depenses = AutreDepense.objects.filter(depense=depense)
    context = {
        'user': user,
        'projet': projet,
        'depense': depense,
        'autre_depenses': autre_depenses
    }
    if request.method == 'POST':
        salaire = request.POST.get('salaire')
        consommable_divers = request.POST.get('consommable_diver')
        equipement_materiel = request.POST.get('equipement_materiel')

        intitules = request.POST.getlist('intitules[]')
        prix = request.POST.getlist('couts[]')


        if(all([salaire])):
            with transaction.atomic():
                depense.salaire_avantages = salaire
                depense.consommable_divers = consommable_divers
                depense.equipement_materiel = equipement_materiel

                depense.save()
        
        if (all([intitules, prix])):
            for name, price in zip(intitules, prix):
                try:
                    AutreDepense.objects.get(
                        depense = depense,
                        intitule = name,
                        prix = price
                    )
                except ObjectDoesNotExist:
                    autre_depense = AutreDepense.objects.create(
                        depense = depense,
                        intitule = name,
                        prix = price
                    )
                    autre_depense.save()

        messages.success(request, "DEPENSES DE FONCTIONNEMENT : modifier avec succes")
        return render(request, 'service/modifier_depense.html', context)
    
    return render(request, 'service/modifier_depense.html', context)

@login_required(login_url='/login/')
@all_user_required
def view_depense(request, projet_id, depense_id):
    user = request.user
    projet = get_object_or_404(Projet, pk=projet_id)
    depenses = get_object_or_404(Depense, pk=depense_id)
    autre_depense = AutreDepense.objects.filter(depense=depenses)
    context = {
        'user': user,
        'projet': projet,
        'depenses': depenses,
        'autre_depense': autre_depense
    }
    return render(request, 'service/view_depense.html', context)

@login_required(login_url='/login/')
@responsable_required
def parametres(request):
    context = {

    }
    return render(request, 'service/parametres.html', context)
